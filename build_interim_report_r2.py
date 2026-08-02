#!/usr/bin/env python3
"""
build_interim_report_r2.py
==========================
產生「期中報告 r2」.docx —— 格式依 F:/SSA/20260729期中報告/期中報告-星文整理-01.doc：

  1. 官方前置頁（Section 1，無頁首）
       封面（MID-RPT-01 / TASA-S-1150268 / 期中進度報告 / 國防大學理工學院）
       簽核頁（Approval / Reviewed / Prepared 四人欄位）
       改版／變更記錄表（中英雙語表頭）
  2. 本文（Section 2，頁首含報告名稱與文件編號）
       目錄（TOC 欄位）→ 由 Word 開啟時自動更新
       PART A–G ＋ Appendix A–E

  字型：標楷體（CJK）／Times New Roman（Latin）；內文 12pt、標題 14pt 粗體
  版面：A4，上下 50.4pt、左右 57.6pt

內容來源：docs/期中報告_MEME_TLE_20260715_r2.md（已修正版）
輸出：    docs/期中報告_MEME_TLE_20260715_r2.docx

用法：python build_interim_report_r2.py
"""
from __future__ import annotations

import re
import sys
from pathlib import Path

# 複用既有 md→docx 轉換機制（表格／圖片／方程式／行內樣式）
import build_interim_report_doc as B

from docx import Document
from docx.enum.section import WD_SECTION
from docx.enum.style import WD_STYLE_TYPE
from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_TAB_ALIGNMENT, WD_BREAK
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Pt, Inches, RGBColor

try:
    sys.stdout.reconfigure(encoding="utf-8")
except Exception:
    pass

ROOT = Path("f:/GitHub/Sat_TraingDataExtension")
DOCS = ROOT / "docs"
SRC_MD = DOCS / "期中報告_MEME_TLE_20260715_r7.md"
OUT = DOCS / "期中報告_MEME_TLE_20260715_r7.docx"

DOC_NO = "MID-RPT-01"
DOC_ISSUE = "Issue 00 Revision 00"
DOC_DATE = "2026/07/23"
PROJ_NO = "TASA-S-1150268"
TITLE1 = "智慧化低軌通訊衛星軌道異常及"
TITLE2 = "太空事件偵測演算法研究"
RPT_KIND = "期中進度報告"
ORG_EN = "CHUNG CHENG INSTITUTE OF TECHNOLOGY, NDU"
ORG_ZH = "國防大學理工學院"

SIGNERS = [
    ("Approval by:", "核可", "李宜珊／計畫主持人"),
    ("Reviewed by:", "審查", "吳志堅／兼任研究員"),
    ("", "", "劉志堅／兼任研究員"),
    ("Prepared by:", "擬定", "方星文／兼任研究員"),
]

# ── 依範本覆寫字型 ───────────────────────────────────────────────────────────
B.CJK = "標楷體"
B.LATIN = "Times New Roman"
INK = B.INK
MUTED = B.MUTED


def _kai(run, size=12, bold=False, color=INK, italic=False):
    run.font.size = Pt(size); run.bold = bold; run.italic = italic
    run.font.color.rgb = RGBColor.from_string(color)
    run.font.name = B.LATIN
    rPr = run._element.get_or_add_rPr()
    rPr.get_or_add_rFonts().set(qn("w:eastAsia"), B.CJK)


# ── 標題純文字化：去除舊編號（18.3／13.1…）與跨章節引用（第X節／§13.x）───────
_LEADNUM_RE = re.compile(r"^\d{1,2}(?:\.\d{1,2})+\s*[　 ]*")
_BOLDLINE_RE = re.compile(r"^(\*\*)(.+?)(\*\*)([:：]?\s*)$")
_XREF_FIXES = [
    (re.compile(r"（回應第[一二三四五六七八九十]+節根因，"), "（"),
    (re.compile(r"；嚴重度定義同第[一二三四五六七八九十]+節："), "；嚴重度定義："),
    (re.compile(r"（對映\s*§\s*[0-9]+(?:\.[0-9]+)?[^）]*）"), ""),
    (re.compile(r"，AUC 取自\s*§\s*[0-9]+(?:\.[0-9]+)?\s*"), "，AUC 取自"),
    (re.compile(r"（第[一二三四五六七八九十]+節）"), ""),
    (re.compile(r"（見第[一二三四五六七八九十]+節[^）]*）"), ""),
]


def _clean_title(t: str, strip_leadnum: bool = True) -> str:
    """標題文字淨化：移除跨章引用；非圖表題註再移除舊數字編號前綴。"""
    for rx, rep in _XREF_FIXES:
        t = rx.sub(rep, t)
    if strip_leadnum and not re.match(r"^[表圖]", t):
        t = _LEADNUM_RE.sub("", t)
    return t


_BOLDNUM_HEAD_RE = re.compile(r"^\*\*\d{1,2}(?:\.\d{1,2})+\s*[　 ]*")


def _sanitize_titles(md: str) -> str:
    """偽標題淨化：整行粗體套標題淨化；行首「**18.1 …」型剝除舊編號。"""
    out = []
    for ln in md.split("\n"):
        s = ln.strip()
        m = _BOLDLINE_RE.match(s)
        if m:
            ln = f"{m.group(1)}{_clean_title(m.group(2))}{m.group(3)}{m.group(4)}"
        elif _BOLDNUM_HEAD_RE.match(s):
            ln = _BOLDNUM_HEAD_RE.sub("**", s)
            for rx, rep in _XREF_FIXES:
                ln = rx.sub(rep, ln)
        out.append(ln)
    return "\n".join(out)


def _heading_kai(doc, text, level):
    """範本樣式：標題掛內建 Heading 樣式（進目錄）；段落間距依範本＝前 18pt／後 6pt。"""
    p = doc.add_paragraph(style=f"Heading {min(level, 3)}")
    pf = p.paragraph_format
    pf.space_before = Pt(18)
    pf.space_after = Pt(6)
    r = p.add_run(_clean_title(text.replace("`", ""), strip_leadnum=False))
    _kai(r, size=14, bold=True, color=INK)
    return p


B._heading = _heading_kai          # 讓 md_to_docx 走範本標題樣式
B._cjk = _kai                      # 內文改標楷體 12pt


# ── 欄位（TOC／頁碼）─────────────────────────────────────────────────────────
def _field(paragraph, instr: str, placeholder: str = ""):
    r = paragraph.add_run()
    fc = OxmlElement("w:fldChar"); fc.set(qn("w:fldCharType"), "begin")
    it = OxmlElement("w:instrText"); it.set(qn("xml:space"), "preserve"); it.text = instr
    sep = OxmlElement("w:fldChar"); sep.set(qn("w:fldCharType"), "separate")
    t = OxmlElement("w:t"); t.text = placeholder
    end = OxmlElement("w:fldChar"); end.set(qn("w:fldCharType"), "end")
    for el in (fc, it, sep, t, end):
        r._r.append(el)
    return r


def _page_footer(section, start_arabic=True):
    f = section.footer
    p = f.paragraphs[0] if f.paragraphs else f.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _field(p, " PAGE ", "1")
    for r in p.runs:
        _kai(r, size=10, color=MUTED)


def _header_body(section):
    h = section.header
    p = h.paragraphs[0] if h.paragraphs else h.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    ts = p.paragraph_format.tab_stops
    _kai(p.add_run(f"{TITLE1}{TITLE2}　　{DOC_NO}"), size=9, color=MUTED)
    # 頁首下框線
    pPr = p._p.get_or_add_pPr()
    pbdr = pPr.makeelement(qn("w:pBdr"), {})
    bottom = pbdr.makeelement(qn("w:bottom"),
                              {qn("w:val"): "single", qn("w:sz"): "4",
                               qn("w:space"): "1", qn("w:color"): "808080"})
    pbdr.append(bottom); pPr.append(pbdr)


def _blank(doc, n=1, size=12):
    for _ in range(n):
        p = doc.add_paragraph()
        _kai(p.add_run(""), size=size)


def _center(doc, text, size, bold=False, latin=False):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(2)
    r = p.add_run(text)
    _kai(r, size=size, bold=bold)
    return p


# ── 前置頁 ───────────────────────────────────────────────────────────────────
def cover(doc):
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(0)
    _kai(p.add_run(DOC_NO), size=12)
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(0)
    _kai(p.add_run(DOC_ISSUE), size=12)
    p = doc.add_paragraph(); p.paragraph_format.space_after = Pt(0)
    _kai(p.add_run(DOC_DATE), size=12)

    _blank(doc, 6)
    _center(doc, TITLE1, 24, True)
    _center(doc, TITLE2, 24, True)
    _blank(doc, 1)
    _center(doc, PROJ_NO, 24, True)
    _blank(doc, 1)
    _center(doc, RPT_KIND, 24, True)
    _blank(doc, 8)
    _center(doc, ORG_EN, 18, True)
    _center(doc, ORG_ZH, 18, True)
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def signature(doc):
    _blank(doc, 2)
    _center(doc, TITLE1, 24, True)
    _center(doc, TITLE2, 24, True)
    _center(doc, RPT_KIND, 24, True)
    _blank(doc, 4)

    for en, zh, who in SIGNERS:
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(2)
        ts = p.paragraph_format.tab_stops
        ts.add_tab_stop(Inches(1.55), WD_TAB_ALIGNMENT.LEFT)
        _kai(p.add_run(f"{en}\t_______________________________"), size=14)

        p2 = doc.add_paragraph()
        p2.paragraph_format.space_after = Pt(16)
        ts2 = p2.paragraph_format.tab_stops
        ts2.add_tab_stop(Inches(1.55), WD_TAB_ALIGNMENT.LEFT)
        ts2.add_tab_stop(Inches(5.6), WD_TAB_ALIGNMENT.LEFT)
        _kai(p2.add_run(f"{zh}\t{who}\tDate"), size=12)

    _blank(doc, 3)
    _center(doc, ORG_EN, 18, True)
    _center(doc, ORG_ZH, 18, True)
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def revision_record(doc):
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    _kai(p.add_run("Revision/Change Record"), size=12, bold=True)
    p = doc.add_paragraph(); p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_after = Pt(8)
    _kai(p.add_run("改版／變更記錄"), size=12, bold=True)

    heads = [("Symbol", "記號"), ("Document Date", "文件日期"),
             ("Authorization Date", "核可日期"),
             ("Revision/Change Description", "改版／變更說明"),
             ("Page Affected", "影響頁次")]
    t = doc.add_table(rows=1, cols=len(heads))
    t.style = "Table Grid"
    for i, (en, zh) in enumerate(heads):
        c = t.rows[0].cells[i]
        c.paragraphs[0].clear(); c.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        _kai(c.paragraphs[0].add_run(en), size=11, bold=True)
        p2 = c.add_paragraph(); p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
        _kai(p2.add_run(zh), size=11, bold=True)

    rows = [
        ("0000", DOC_DATE, "", "期中進度報告發行", "全", ),
    ]
    for r in rows:
        cells = t.add_row().cells
        for i, v in enumerate(r):
            cells[i].paragraphs[0].clear()
            if i != 3:
                cells[i].paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
            _kai(cells[i].paragraphs[0].add_run(v), size=10.5)
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def _ensure_caption_styles(doc):
    """建立「圖題」「表題」段落樣式（供圖目錄／表目錄之 TOC \\t 欄位擷取）。"""
    for name in ("圖題", "表題"):
        if name not in [s.name for s in doc.styles]:
            st = doc.styles.add_style(name, WD_STYLE_TYPE.PARAGRAPH)
            st.base_style = doc.styles["Normal"]
            st.font.name = B.LATIN
            st.font.size = Pt(9.5)
            st.element.rPr.rFonts.set(qn("w:eastAsia"), B.CJK)
            st.quick_style = True


_FIG_CAP = re.compile(r"^圖[\s　]+[0-9A-Za-zF]")
_TAB_CAP = re.compile(r"^表[\s　]+[0-9A-Za-zF]")

# r6.md 圖說內嵌於 ![圖 N　說明](path) 之 alt；B._image 不渲染 alt。
# 轉為「圖片 + 其後 blockquote 圖說」，使圖說成為可見段落並供圖目錄擷取。
_FIGIMG_RE = re.compile(r"^!\[(圖[^\]]*)\]\(([^)]*)\)\s*$", re.M)


def _inject_fig_captions(md: str) -> str:
    return _FIGIMG_RE.sub(lambda m: f"![]({m.group(2)})\n\n> {m.group(1)}", md)


def _shade_para(p, fill: str):
    """段落底色（w:shd）。"""
    pPr = p._p.get_or_add_pPr()
    for old in pPr.findall(qn("w:shd")):
        pPr.remove(old)
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear"); shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), fill)
    pPr.append(shd)


TAB_CAP_FILL = "D9E2F3"   # 表題淺藍底（仿範本截圖）


def _tag_captions(doc):
    """把圖題／表題段落掛上專屬樣式（保留置中外觀），供 TOF 欄位擷取；表題加淺藍底。"""
    nfig = ntab = 0
    for p in doc.paragraphs:
        sn = p.style.name if p.style else ""
        if sn.startswith("Heading") or sn.startswith("List") or sn in ("圖題", "表題"):
            continue
        t = p.text.strip()
        if _FIG_CAP.match(t):
            p.style = doc.styles["圖題"]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER; nfig += 1
        elif _TAB_CAP.match(t):
            p.style = doc.styles["表題"]; p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            _shade_para(p, TAB_CAP_FILL); ntab += 1
    return nfig, ntab


# ── 內文圖／表引用黃底凸顯 ───────────────────────────────────────────────────
_REF_RE = re.compile(r"[表圖][ 　]?(?:\d{1,2}|[A-G])(?:[-–－][0-9A-Za-z]{1,4})?")
_SKIP_STYLES = ("圖題", "表題")


def _hl_yellow(run):
    rPr = run._element.get_or_add_rPr()
    for old in rPr.findall(qn("w:highlight")):
        rPr.remove(old)
    hl = OxmlElement("w:highlight"); hl.set(qn("w:val"), "yellow")
    rPr.append(hl)


def _highlight_refs_in_para(p) -> int:
    """把段內「表 N／圖 N」引用切成獨立 run 並上黃底；回傳命中數。"""
    from copy import deepcopy
    n = 0
    for r in list(p.runs):
        text = r.text
        if not text or not _REF_RE.search(text):
            continue
        rpr0 = r._r.find(qn("w:rPr"))
        rpr_tpl = deepcopy(rpr0) if rpr0 is not None else None
        segs, pos = [], 0
        for m in _REF_RE.finditer(text):
            if m.start() > pos:
                segs.append((text[pos:m.start()], False))
            segs.append((m.group(), True)); pos = m.end()
        if pos < len(text):
            segs.append((text[pos:], False))
        # 首段沿用原 run，其餘依原格式複製新 run 插於其後
        r.text = segs[0][0]
        if segs[0][1]:
            _hl_yellow(r); n += 1
        anchor = r._r
        for seg, ism in segs[1:]:
            nr = OxmlElement("w:r")
            if rpr_tpl is not None:
                nr.append(deepcopy(rpr_tpl))
            t = OxmlElement("w:t"); t.set(qn("xml:space"), "preserve"); t.text = seg
            nr.append(t)
            anchor.addnext(nr); anchor = nr
            if ism:
                from docx.text.run import Run
                _hl_yellow(Run(nr, p)); n += 1
    return n


def _highlight_refs(doc) -> int:
    """全文（含表格內）圖／表引用黃底；跳過標題與圖題／表題本身。"""
    total = 0
    def eligible(p):
        sn = p.style.name if p.style else ""
        return not (sn.startswith("Heading") or sn in _SKIP_STYLES)
    for p in doc.paragraphs:
        if eligible(p):
            total += _highlight_refs_in_para(p)
    for tbl in doc.tables:
        for row in tbl.rows:
            for cell in row.cells:
                for p in cell.paragraphs:
                    if eligible(p):
                        total += _highlight_refs_in_para(p)
    return total


def _tof(doc, title, style_name):
    """插入一個以段落樣式擷取的目錄（圖目錄／表目錄）。"""
    _center(doc, title, 16, True)
    _blank(doc, 1)
    p = doc.add_paragraph()
    _field(p, f' TOC \\h \\z \\t "{style_name},1" ', "（將於 Word 更新後產生）")
    for r in p.runs:
        _kai(r, size=11, color=MUTED)
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)


def toc(doc):
    _center(doc, "目錄", 16, True)
    _blank(doc, 1)
    p = doc.add_paragraph()
    _field(p, r' TOC \o "1-3" \h \z \u ', "（請在 Word 中按 Ctrl+A 後 F9 更新目錄）")
    for r in p.runs:
        _kai(r, size=11, color=MUTED)
    doc.add_paragraph().add_run().add_break(WD_BREAK.PAGE)
    # 目錄之後：圖目錄、表目錄（依樣式擷取；標題文字仿範本）
    _tof(doc, "圖目錄", "圖題")
    _tof(doc, "表目錄", "表題")


# ── PART 重構映射 ────────────────────────────────────────────────────────────
# (原 md 標題文字前綴, 新編號, 新標題)；level 由原始 md 決定（H3→H2、H4→H3）
PARTS = [
    ("PART A : 簡介", [
        ("導讀（撰寫體例說明）", "A.1", "撰寫體例與閱讀指引"),
    ]),
    ("PART B : 相關資訊", [
        ("摘要（結論先行）", "B.1", "摘要（結論先行）"),
        ("送審重點一頁摘要", "B.2", "送審重點一頁摘要：真值定義、驗收主指標、不適用場景"),
        ("研究意義與應用價值（四大價值）", "B.3", "研究意義與應用價值（四大價值）"),
        ("一、為什麼是現在", "B.4", "研究背景與動機"),
    ]),
    ("PART C : Layer 1 — TLE 差分機動偵測", [
        ("二、不是一張圖", "C.1", "三層架構總覽"),
        ("2.1", "C.2", "TLE 欄位與差分定義"),
        ("2.2", "C.3", "J2 攝動修正"),
        ("2.3", "C.4", "P1–P6 規則與觸發條件"),
        ("2.4", "C.5", "Layer 1 全量驗證結果"),
    ]),
    ("PART D : Layer 2 — 統計變點偵測與融合評分器", [
        ("三、把統計學帶回軌道判讀", "D.1", "Layer 2 統計異常偵測器"),
        ("四、融合評分器", "D.2", "融合評分器：五通道證據整合"),
    ]),
    ("PART E : Layer 3 — 機器學習、物理交叉驗證與路由", [
        ("五、機器學習的價值與邊界", "E.1", "機器學習的價值與邊界"),
        ("六、物理作為第三方證人", "E.2", "NRLMSIS 阻力殘差：物理第三方驗證"),
        ("七、再入守門", "E.3", "再入守門：知道什麼時候不該用模型"),
        ("八、不押注單一方法", "E.4", "路由策略：不押注單一方法"),
    ]),
    ("PART F : 能力驗證、系統應用與延伸專題", [
        ("九、從單星到星系", "F.1", "星系級異常分析"),
        ("十、知道極限在哪", "F.2", "偵測下限的量化"),
        ("十一、機動時刻偵測", "F.3", "機動時刻偵測（burn epoch）"),
        ("十二、系統應用", "F.4", "系統應用：分析環境而非黑盒警報器"),
        ("十三、驗證與訓練策略", "F.5", "驗證與訓練策略：泛化落差如何收斂"),
        ("13.1", "F.5.1", "全庫實測背書：1,000 顆隨機酬載三方法比較"),
        ("13.2", "F.5.2", "三層同一擂台補充驗證：相同測試集／GT／單元"),
        ("13.3", "F.5.3", "三層交叉分析：L3 修正 L1／L2 之漏警與誤警"),
        ("13.4", "F.5.4", "L3 泛化穩定度：增益非過擬合特定分布之驗證"),
        ("13.5", "F.5.5", "評估口徑對照：指標↔評估單位↔操作目標"),
        ("十四、可重複性", "F.6", "可重複性：把判讀變成可審查的流程"),
        ("十五、新功能測試", "F.7", "在軌釋放事件的重建（RPO）"),
        ("十六、FORMOSAT 系列專章", "F.8", "FORMOSAT 系列專章"),
        ("16.1", "F.8.1", "福衛三：退役星系是物理模型最素淨的考場"),
        ("16.2", "F.8.2", "福衛五／獵風者／福衛八A：三種「正常」的樣子"),
        ("16.3", "F.8.3", "福衛七星群：同步機動作業、一顆掉隊、與軌道面幾何"),
        ("十七、延伸專題", "F.9", "延伸專題：Starlink 訊號替代 GPS 之定位"),
        ("17.1", "F.9.1", "兩條路線的實驗流程"),
        ("17.2", "F.9.2", "定位天花板的量化：2,600 萬點實測"),
        ("17.3", "F.9.3", "三項創見與建議（非契約範圍之外溢效益討論）"),
    ]),
    ("PART G : 期中成果彙整、期末待辦與總結", [
        ("十八、期中成果彙整", "G.1", "期中成果彙整、期末待辦與建議事項"),
        ("十九、Starlink 進入台灣市場", "G.2", "Starlink 進入台灣市場之研究情境與本專案意涵"),
        ("結語", "G.3", "結語"),
    ]),
]

APPENDIX = [
    ("附錄 A：指標速查表", "Appendix A", "指標速查表"),
    ("附錄 B：關鍵程式與資料對照", "Appendix B", "關鍵程式與資料對照"),
    ("附錄 C：參考文獻", "Appendix C", "參考文獻"),
    ("附錄 D：名詞與縮寫對照表", "Appendix D", "名詞與縮寫對照表"),
    ("附錄 E：全文關鍵參數速查表", "Appendix E", "全文關鍵參數速查表"),
    ("附錄 F：系統統整", "Appendix F", "系統統整——架構責任鏈、真值、泛化與部署"),
    ("附錄 G：Layer 2 四偵測器", "Appendix G", "Layer 2 四偵測器之數學定義"),
]

_H_RE = re.compile(r"^(#{1,6})\s+(.*)$")


def parse_blocks(md: str):
    """切成 [(level, title, body_text), ...]；level 0 = 前言（無標題）。"""
    lines = md.split("\n")
    blocks, cur = [], (0, "", [])
    for ln in lines:
        m = _H_RE.match(ln)
        if m:
            blocks.append((cur[0], cur[1], "\n".join(cur[2])))
            cur = (len(m.group(1)), m.group(2).strip(), [])
        else:
            cur[2].append(ln)
    blocks.append((cur[0], cur[1], "\n".join(cur[2])))
    return blocks


def find_block(blocks, prefix):
    for i, (lv, title, body) in enumerate(blocks):
        if title.startswith(prefix) or prefix in title[:len(prefix) + 4]:
            return i
    return None


def build():
    md = SRC_MD.read_text(encoding="utf-8")
    blocks = parse_blocks(md)
    used = set()

    doc = Document()

    # 版面：A4、範本邊界
    for s in doc.sections:
        s.page_width, s.page_height = Inches(8.27), Inches(11.69)
        s.top_margin = s.bottom_margin = Pt(50.4)
        s.left_margin = s.right_margin = Pt(57.6)

    # 內文預設樣式（仿範本：標楷體 12pt、單行距、段後 6pt）
    st = doc.styles["Normal"]
    st.font.name = B.LATIN; st.font.size = Pt(12)
    st.element.rPr.rFonts.set(qn("w:eastAsia"), B.CJK)
    st.paragraph_format.space_after = Pt(6)
    st.paragraph_format.line_spacing = 1.0

    # ── Section 1：前置頁（仿範本：首頁不同＝封面無頁碼，其後頁尾置中頁碼）──
    front_sec = doc.sections[0]
    sectPr = front_sec._sectPr
    if sectPr.find(qn("w:titlePg")) is None:
        sectPr.append(OxmlElement("w:titlePg"))
    _page_footer(front_sec)
    cover(doc)
    signature(doc)
    revision_record(doc)

    # ── Section 2：本文（仿範本：頁首空白、頁尾置中頁碼）──
    body_sec = doc.add_section(WD_SECTION.NEW_PAGE)
    body_sec.page_width, body_sec.page_height = Inches(8.27), Inches(11.69)
    body_sec.top_margin = body_sec.bottom_margin = Pt(50.4)
    body_sec.left_margin = body_sec.right_margin = Pt(57.6)
    _page_footer(body_sec)

    _ensure_caption_styles(doc)
    toc(doc)

    n_mapped = 0
    for part_title, items in PARTS:
        _heading_kai(doc, part_title, 1)
        for prefix, num, new_title in items:
            idx = find_block(blocks, prefix)
            if idx is None:
                print(f"  [warn] 找不到區塊：{prefix}")
                continue
            lv, title, body = blocks[idx]
            used.add(idx)
            level = 2 if num.count(".") == 1 else 3
            _heading_kai(doc, f"{num}　{new_title}", level)
            if body.strip():
                B.md_to_docx(doc, _inject_fig_captions(_sanitize_titles(body)), DOCS)
            n_mapped += 1

    for prefix, num, new_title in APPENDIX:
        idx = find_block(blocks, prefix)
        if idx is None:
            print(f"  [warn] 找不到附錄：{prefix}")
            continue
        used.add(idx)
        _heading_kai(doc, f"{num}：{new_title}", 1)
        if blocks[idx][2].strip():
            B.md_to_docx(doc, _inject_fig_captions(_sanitize_titles(blocks[idx][2])), DOCS)
        n_mapped += 1

    # 未被映射到的區塊（標題頁等）如實回報
    skipped = [blocks[i][1] for i in range(len(blocks))
               if i not in used and blocks[i][1] and blocks[i][0] <= 4]
    if skipped:
        print(f"  [info] 未納入本文之原標題（共 {len(skipped)}）：{skipped}")

    nfig, ntab = _tag_captions(doc)
    print(f"  [caption] 掛樣式：圖題 {nfig}、表題 {ntab}（表題淺藍底；供圖/表目錄擷取）")

    nhl = _highlight_refs(doc)
    print(f"  [highlight] 內文圖／表引用黃底：{nhl} 處")

    doc.save(str(OUT))
    print(f"[docx] {OUT}  （映射 {n_mapped} 個區塊）")
    return OUT


def verify(path: Path):
    """輸出前逐項檢查（規則 1–5）。"""
    d = Document(str(path))
    probs = []

    # 1. 標題純文字：不得殘留舊數字編號或「第X節」引用
    bad_hd = []
    for p in d.paragraphs:
        sn = p.style.name if p.style else ""
        t = p.text.strip()
        if sn.startswith("Heading") and t:
            if re.match(r"^\d{1,2}\.\d", t) or re.search(r"第[一二三四五六七八九十]+節", t) or "§" in t:
                bad_hd.append(t[:40])
    # 偽標題殘留（粗體舊編號）
    ghosts = [p.text[:40] for p in d.paragraphs
              if re.match(r"^\d{1,2}\.\d{1,2}\s", p.text.strip())
              and p.runs and p.runs[0].bold]
    if bad_hd: probs.append(f"標題殘留舊標號/跨章引用：{bad_hd}")
    if ghosts: probs.append(f"粗體偽標題殘留：{ghosts}")

    # 2/3. 頁首空白、頁尾頁碼、標題間距 18/6
    hd_txt = [pa.text for s in d.sections for pa in s.header.paragraphs if pa.text.strip()]
    if hd_txt: probs.append(f"頁首非空白：{hd_txt}")
    ft_ok = any("PAGE" in s.footer._element.xml for s in d.sections)
    if not ft_ok: probs.append("頁尾缺頁碼欄位")
    sp_bad = [p.text[:25] for p in d.paragraphs
              if p.style.name.startswith("Heading") and p.text.strip()
              and (p.paragraph_format.space_before is None
                   or round(p.paragraph_format.space_before.pt) != 18
                   or round(p.paragraph_format.space_after.pt) != 6)]
    if sp_bad: probs.append(f"標題間距非 18/6pt：{sp_bad[:3]}")

    # 4. 黃底引用計數
    xml = d.element.xml
    nhl = xml.count('w:val="yellow"')
    # 5. 表題藍底計數
    nshd = xml.count(f'w:fill="{TAB_CAP_FILL}"')

    # 目錄三件套
    n_toc = xml.count("TOC")
    print(f"  [verify] 黃底引用 {nhl} 處｜表題藍底 {nshd} 段｜TOC 欄位 {n_toc} 個")
    if probs:
        for x in probs: print(f"  [verify][FAIL] {x}")
    else:
        print("  [verify] 規則 1–5 檢查全數通過")
    return not probs


def update_fields(path: Path):
    """用 Word COM 更新目錄欄位並重新存檔。"""
    try:
        import win32com.client as win32
        # 用動態繫結，避開 gencache 快取損壞（CLSIDToClassMap 缺失）
        app = win32.Dispatch("Word.Application")
        app.Visible = False
        d = app.Documents.Open(str(path.resolve()))
        d.Fields.Update()
        for toc_ in d.TablesOfContents:
            toc_.Update()
        d.Repaginate()
        pages = d.ComputeStatistics(2)
        d.Save(); d.Close(False); app.Quit()
        print(f"[word] 目錄已更新，共 {pages} 頁")
    except Exception as e:
        print(f"[word] 目錄更新略過（請在 Word 中按 Ctrl+A→F9）：{e}")


if __name__ == "__main__":
    p = build()
    verify(p)
    update_fields(p)
    print("完成。")
