"""md_to_docx.py
將論文一與論文二的 Markdown 轉換為 .docx（Word）檔案。
只依賴標準函式庫 + python-docx，無需 pandoc。

執行：python docs/md_to_docx.py
輸出：docs/paper1_tle_maneuver_detection_zh.docx
       docs/paper2_lightgbm_classifier_zh.docx
"""
import re
import os
from pathlib import Path
from docx import Document
from docx.shared import Pt, Cm, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

DOCS_DIR = Path(__file__).parent

# ── LaTeX 數式轉可讀文字 ──────────────────────────────────────────────────────
_LATEX_MAP = [
    (r"\\dot\{\\Omega\}",   "Ω̇"),
    (r"\\dot\{\\omega\}",   "ω̇"),
    (r"\\Delta",            "Δ"),
    (r"\\Omega",            "Ω"),
    (r"\\omega",            "ω"),
    (r"\\theta",            "θ"),
    (r"\\beta",             "β"),
    (r"\\alpha",            "α"),
    (r"\\sigma",            "σ"),
    (r"\\tau",              "τ"),
    (r"\\mu",               "μ"),
    (r"\\pi",               "π"),
    (r"\\geq",              "≥"),
    (r"\\leq",              "≤"),
    (r"\\neq",              "≠"),
    (r"\\approx",           "≈"),
    (r"\\times",            "×"),
    (r"\\cdot",             "·"),
    (r"\\pm",               "±"),
    (r"\\infty",            "∞"),
    (r"\\sqrt\{([^}]+)\}",  r"√(\1)"),
    (r"\\frac\{([^}]+)\}\{([^}]+)\}", r"\1/\2"),
    (r"\\left\(",           "("),
    (r"\\right\)",          ")"),
    (r"\\left\[",           "["),
    (r"\\right\]",          "]"),
    (r"\\text\{([^}]+)\}",  r"\1"),
    (r"\\mathrm\{([^}]+)\}",r"\1"),
    (r"\\mathbf\{([^}]+)\}",r"\1"),
    (r"\\begin\{cases\}",   "{ "),
    (r"\\end\{cases\}",     ""),
    (r"\\\\",               "  /  "),
    (r"\^{([^}]+)}",        r"^\1"),   # superscript
    (r"\^([A-Za-z0-9*])",   r"^\1"),
    (r"_\{([^}]+)\}",       r"_\1"),   # subscript
    (r"_([A-Za-z0-9])",     r"_\1"),
    (r"J_2",                "J₂"),
    (r"R_E",                "Rₑ"),
    (r"\$",                 ""),
]

def latex_to_text(s: str) -> str:
    for pat, rep in _LATEX_MAP:
        s = re.sub(pat, rep, s)
    # strip remaining backslash commands
    s = re.sub(r"\\[a-zA-Z]+", "", s)
    s = re.sub(r"[{}]", "", s)
    return s.strip()

def clean_math(s: str) -> str:
    """Strip display/inline math delimiters and convert LaTeX to text."""
    # Display math $$...$$
    s = re.sub(r"\$\$(.+?)\$\$", lambda m: "[" + latex_to_text(m.group(1)) + "]",
               s, flags=re.DOTALL)
    # Inline math $...$
    s = re.sub(r"\$(.+?)\$", lambda m: latex_to_text(m.group(1)), s)
    return s


# ── 行內樣式解析 ──────────────────────────────────────────────────────────────
def add_inline_runs(para, text: str):
    """解析 **bold**、*italic*、`code`，並依序加入 Run。"""
    # 合法字元範圍
    pattern = re.compile(
        r"\*\*(.+?)\*\*"        # **bold**
        r"|\*(.+?)\*"           # *italic*
        r"|`(.+?)`"             # `code`
        r"|\!\[([^\]]*)\]\([^)]*\)"  # ![img](path) → skip in inline
        r"|(\[([^\]]+)\]\([^)]*\))"  # [link text](url) → text only
        r"|(.+?)"               # plain text
        r"|(\s+)",              # whitespace
        re.DOTALL)

    for m in pattern.finditer(text):
        bold_t, ital_t, code_t, img_t, _, link_t, plain_t, ws_t = m.groups()
        if bold_t:
            run = para.add_run(bold_t)
            run.bold = True
        elif ital_t:
            run = para.add_run(ital_t)
            run.italic = True
        elif code_t:
            run = para.add_run(code_t)
            run.font.name = "Courier New"
            run.font.size = Pt(9)
        elif link_t:
            para.add_run(link_t)
        elif plain_t:
            para.add_run(plain_t)
        elif ws_t:
            para.add_run(ws_t)


def add_para_text(para, raw: str):
    """先清理 math、再寫入行內樣式。"""
    raw = clean_math(raw)
    # 清理 Markdown 轉義
    raw = raw.replace(r"\*", "*").replace(r"\_", "_").replace(r"\|", "|")
    add_inline_runs(para, raw)


# ── 表格處理 ──────────────────────────────────────────────────────────────────
def strip_cell(s: str) -> str:
    return clean_math(s.strip().strip("*").strip())

def is_separator_row(row_text: str) -> bool:
    return bool(re.match(r"^\|?[\s\-:]+(\|[\s\-:]+)*\|?$", row_text.strip()))

def parse_table_row(line: str):
    parts = [c.strip() for c in line.strip().strip("|").split("|")]
    return parts

def build_word_table(doc, rows):
    """rows: list of list[str] — 第一行為表頭。"""
    if not rows:
        return
    ncols = max(len(r) for r in rows)
    tbl = doc.add_table(rows=len(rows), cols=ncols)
    tbl.style = "Table Grid"
    tbl.alignment = WD_TABLE_ALIGNMENT.CENTER
    for ri, row in enumerate(rows):
        for ci, cell_text in enumerate(row):
            if ci >= ncols:
                break
            cell = tbl.cell(ri, ci)
            cell.text = ""
            p = cell.paragraphs[0]
            p.clear()
            run = p.add_run(strip_cell(cell_text))
            if ri == 0:
                run.bold = True
            p.paragraph_format.space_before = Pt(2)
            p.paragraph_format.space_after  = Pt(2)
    doc.add_paragraph()  # 表格後空行


# ── 圖片插入 ──────────────────────────────────────────────────────────────────
def try_insert_image(doc, img_path_str: str):
    img_path = DOCS_DIR / img_path_str
    if img_path.exists():
        try:
            p = doc.add_paragraph()
            p.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run = p.add_run()
            run.add_picture(str(img_path), width=Inches(5.8))
        except Exception as e:
            doc.add_paragraph(f"[圖片：{img_path_str}  — 插入失敗：{e}]")
    else:
        doc.add_paragraph(f"[圖片：{img_path_str}  — 檔案未找到]")


# ── 文件樣式設定 ──────────────────────────────────────────────────────────────
def setup_doc_styles(doc: Document):
    """設定中文字體與基本樣式。"""
    from docx.shared import Pt, RGBColor
    from docx.oxml.ns import qn

    style_normal = doc.styles["Normal"]
    style_normal.font.name = "Microsoft YaHei"
    style_normal.font.size = Pt(11)
    style_normal._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    for heading_name, size, bold_flag in [
        ("Heading 1", 18, True),
        ("Heading 2", 15, True),
        ("Heading 3", 13, True),
    ]:
        sty = doc.styles[heading_name]
        sty.font.name = "Microsoft YaHei"
        sty.font.size = Pt(size)
        sty.font.bold = bold_flag
        sty.font.color.rgb = RGBColor(0x1a, 0x3a, 0x5c)
        sty._element.rPr.rFonts.set(qn("w:eastAsia"), "Microsoft YaHei")

    # Quote style
    if "Quote" not in [s.name for s in doc.styles]:
        pass
    try:
        q = doc.styles["Quote"]
        q.font.name  = "Microsoft YaHei"
        q.font.size  = Pt(10)
        q.font.italic = True
        q.paragraph_format.left_indent  = Cm(1)
        q.paragraph_format.right_indent = Cm(1)
    except Exception:
        pass


# ── 主轉換邏輯 ────────────────────────────────────────────────────────────────
def convert_md_to_docx(md_path: Path, out_path: Path):
    doc = Document()
    setup_doc_styles(doc)

    # 頁面邊距
    for section in doc.sections:
        section.top_margin    = Cm(2.5)
        section.bottom_margin = Cm(2.5)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(2.5)

    lines = md_path.read_text(encoding="utf-8").splitlines()

    state     = "normal"   # normal | table | code_block
    tbl_rows  = []
    code_lines = []

    i = 0
    while i < len(lines):
        raw = lines[i]
        stripped = raw.strip()

        # ── 程式碼區塊 ──────────────────────────────────────────────
        if stripped.startswith("```"):
            if state != "code_block":
                state = "code_block"
                code_lines = []
                i += 1
                continue
            else:
                # 結束 code block
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(1)
                run = p.add_run("\n".join(code_lines))
                run.font.name = "Courier New"
                run.font.size = Pt(8.5)
                state = "normal"
                code_lines = []
                i += 1
                continue

        if state == "code_block":
            code_lines.append(raw)
            i += 1
            continue

        # ── 表格 ────────────────────────────────────────────────────
        if stripped.startswith("|") or (state == "table" and stripped.startswith("|")):
            if is_separator_row(stripped):
                i += 1
                continue
            if state != "table":
                state = "table"
                tbl_rows = []
            tbl_rows.append(parse_table_row(stripped))
            # 看下一行是否還是表格
            next_line = lines[i+1].strip() if i+1 < len(lines) else ""
            if not next_line.startswith("|") and not is_separator_row(next_line):
                build_word_table(doc, tbl_rows)
                state = "normal"
                tbl_rows = []
            i += 1
            continue

        if state == "table":
            build_word_table(doc, tbl_rows)
            state = "normal"
            tbl_rows = []
            # 不 continue，繼續處理當前行

        # ── 圖片 ─────────────────────────────────────────────────────
        img_match = re.match(r"!\[([^\]]*)\]\(([^)]+)\)", stripped)
        if img_match:
            alt_text = img_match.group(1)
            img_file = img_match.group(2)
            try_insert_image(doc, img_file)
            # 加圖說
            p_cap = doc.add_paragraph()
            p_cap.alignment = WD_ALIGN_PARAGRAPH.CENTER
            run_cap = p_cap.add_run(alt_text if alt_text else img_file)
            run_cap.italic = True
            run_cap.font.size = Pt(9.5)
            i += 1
            continue

        # ── blockquote ─────────────────────────────────────────────
        if stripped.startswith("> "):
            content = stripped[2:]
            # 多行 blockquote 合併
            while i+1 < len(lines) and lines[i+1].strip().startswith("> "):
                i += 1
                content += " " + lines[i].strip()[2:]
            content = clean_math(content)
            try:
                p = doc.add_paragraph(style="Quote")
            except Exception:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(1)
            add_inline_runs(p, content)
            i += 1
            continue

        # ── 標題 ─────────────────────────────────────────────────────
        h_match = re.match(r"^(#{1,4})\s+(.+)", stripped)
        if h_match:
            level = len(h_match.group(1))
            title_text = clean_math(h_match.group(2))
            style_map = {1: "Heading 1", 2: "Heading 2",
                         3: "Heading 3", 4: "Heading 4"}
            style_name = style_map.get(level, "Heading 3")
            p = doc.add_paragraph(style=style_name)
            p.clear()
            add_inline_runs(p, title_text)
            i += 1
            continue

        # ── 水平線 ─────────────────────────────────────────────────
        if stripped in ("---", "***", "___"):
            doc.add_paragraph()  # 空行代替分割線
            i += 1
            continue

        # ── 有序/無序清單 ───────────────────────────────────────────
        li_match = re.match(r"^(\s*)(\d+\.|-|\*)\s+(.+)", raw)
        if li_match:
            indent_len = len(li_match.group(1))
            marker     = li_match.group(2)
            content    = li_match.group(3)
            is_ordered = bool(re.match(r"\d+\.", marker))
            style_name = "List Number" if is_ordered else "List Bullet"
            try:
                p = doc.add_paragraph(style=style_name)
            except Exception:
                p = doc.add_paragraph()
                p.paragraph_format.left_indent = Cm(0.75 + indent_len * 0.25)
            p.clear()
            if is_ordered:
                p = doc.add_paragraph(style="List Number")
                p.clear()
            add_para_text(p, content)
            i += 1
            continue

        # ── 空行 ─────────────────────────────────────────────────────
        if not stripped:
            doc.add_paragraph()
            i += 1
            continue

        # ── 一般段落 ─────────────────────────────────────────────────
        p = doc.add_paragraph()
        p.paragraph_format.space_after = Pt(4)
        add_para_text(p, stripped)
        i += 1

    doc.save(str(out_path))
    size_kb = out_path.stat().st_size // 1024
    print(f"  saved {out_path}  ({size_kb} KB)")


# ─────────────────────────────────────────────────────────────────────────────
if __name__ == "__main__":
    print("轉換中...")
    convert_md_to_docx(
        DOCS_DIR / "paper1_tle_maneuver_detection_zh.md",
        DOCS_DIR / "paper1_tle_maneuver_detection_zh.docx",
    )
    convert_md_to_docx(
        DOCS_DIR / "paper2_lightgbm_classifier_zh.md",
        DOCS_DIR / "paper2_lightgbm_classifier_zh.docx",
    )
    print("完成。")
