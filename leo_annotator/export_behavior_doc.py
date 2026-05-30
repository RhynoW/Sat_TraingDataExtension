"""
export_behavior_doc.py
======================
將 analyze_maneuver_behavior.py 架構說明輸出為 Word 文件。
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import qn
from docx.oxml import OxmlElement
import copy

OUT_PATH = r"output/analyze_maneuver_behavior_py_20260525架構說明.docx"


# ── 格式輔助 ─────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    p.runs[0].font.color.rgb = RGBColor(0x1F, 0x49, 0x7D) if level == 1 else RGBColor(0x2E, 0x74, 0xB5)
    return p


def add_body(doc, text):
    p = doc.add_paragraph(text)
    p.runs[0].font.size = Pt(11)
    return p


def add_code(doc, text):
    """灰底等寬字型的 code block"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    # 段落底色
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F2F2F2")
    pPr.append(shd)
    return p


def add_table(doc, headers, rows, col_widths=None, header_color="2E74B5"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    # 標題列
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = h
        run = cell.paragraphs[0].runs[0]
        run.bold = True
        run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
        run.font.size = Pt(10)
        set_cell_bg(cell, header_color)
    # 資料列
    for ri, row in enumerate(rows):
        tr = table.rows[ri + 1]
        bg = "F2F2F2" if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row):
            cell = tr.cells[ci]
            cell.text = str(val)
            cell.paragraphs[0].runs[0].font.size = Pt(10)
            if bg == "F2F2F2":
                set_cell_bg(cell, bg)
    # 欄寬
    if col_widths:
        for col_i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[col_i].width = Cm(w)
    return table


# ── 主文件 ───────────────────────────────────────────────────────────────────

def build_doc():
    doc = Document()

    # 頁面邊距
    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── 封面 ─────────────────────────────────────────────────────────────────
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("analyze_maneuver_behavior.py")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = sub.add_run("架構說明文件")
    run2.font.size = Pt(16)
    run2.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = meta.add_run("版本：2026-05-25　　分析資料期間：2026-05-01 ～ 2026-05-22")
    run3.font.size = Pt(10)
    run3.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    doc.add_paragraph()
    doc.add_paragraph("─" * 72).runs[0].font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    doc.add_paragraph()

    # ── 0. 目的 ──────────────────────────────────────────────────────────────
    add_heading(doc, "0. 文件目的", level=1)
    add_body(doc,
        "本文件說明 analyze_maneuver_behavior.py 腳本的完整運作邏輯，"
        "包含資料流、物理量計算方法、行為分類決策樹、輸出格式，以及已知限制。"
        "腳本的主要任務是對 validation_full.csv 中被偵測到軌道機動的 1,323 顆 LEO 衛星，"
        "重新抓取逐筆 TLE 轉移資料，分析機動行為特性並分類，最終輸出 Markdown 報告與 CSV。")

    # ── 1. 整體架構 ──────────────────────────────────────────────────────────
    add_heading(doc, "1. 整體架構", level=1)
    add_body(doc, "腳本為線性管線，五個階段依序執行：")
    doc.add_paragraph()
    add_code(doc,
"validation_full.csv          annotations_leo_full.csv\n"
"      ↓ 篩出 maneuver=True           ↓ 取 propulsion_class\n"
"      └──────────── 1,323 顆 NORAD ──────────────┘\n"
"                          ↓\n"
"              DuckDB 查詢逐筆 TLE（62,005 筆）\n"
"                          ↓\n"
"              detect_transitions()  ← 每對相鄰 TLE 計算軌道元素差\n"
"                          ↓\n"
"              classify_behavior()   ← 規則決策樹分類\n"
"                          ↓\n"
"        maneuver_behavior_classified.csv\n"
"        maneuver_behavior_report.md")
    doc.add_paragraph()
    add_table(doc,
        ["階段", "對應函式 / 來源", "說明"],
        [
            ["①", "main() — 載入 CSV", "讀 validation_full.csv，篩選 maneuver_detected=True"],
            ["②", "main() — 讀標註", "讀 annotations_leo_full.csv，以 norad_id 為索引"],
            ["③", "main() — DuckDB", "一次性查詢 1,323 顆衛星在觀測期間的所有 TLE"],
            ["④", "detect_transitions()", "逐對相鄰 TLE 計算 Δa, Δi, Δe, ΔRAAN_res"],
            ["⑤", "classify_behavior()", "規則決策樹判定行為類別"],
            ["⑥", "generate_report()", "輸出 CSV 與 Markdown 報告"],
        ],
        col_widths=[1.2, 4.5, 9.0])

    # ── 2. 物理量計算 ─────────────────────────────────────────────────────────
    add_heading(doc, "2. 物理量計算：detect_transitions()", level=1)
    add_body(doc,
        "對單顆衛星的所有 TLE 依時間排序後，對每對相鄰 TLE（epoch[i-1] → epoch[i]）"
        "計算四個軌道差分，並標記是否超過偵測閾值。間隔超過 7 天的 TLE 對跳過，避免長期缺測造成誤差。")
    doc.add_paragraph()

    add_heading(doc, "2.1 四個軌道差分", level=2)
    add_table(doc,
        ["變數", "計算公式", "物理意義"],
        [
            ["da",       "sma_km[i] − sma_km[i−1]",                    "半長軸差（正值 = 升軌，負值 = 降軌）"],
            ["di",       "inclination[i] − inclination[i−1]",           "傾角差"],
            ["de",       "eccentricity[i] − eccentricity[i−1]",         "離心率差"],
            ["draan_res","(RAAN_raw − RAAN_J2×Δt)",                      "RAAN 差扣除 J2 自然進動後的殘差"],
        ],
        col_widths=[2.0, 6.5, 6.5])
    doc.add_paragraph()

    add_heading(doc, "2.2 J2 RAAN 進動修正", level=2)
    add_body(doc,
        "地球 J2 扁率攝動使 RAAN 產生自然漂移，必須扣除才能識別主動機動。"
        "修正公式（j2_raan_rate 函式，程式碼第 40 行）：")
    add_code(doc,
"n = sqrt(MU / sma³)          # 平均角速度 (rad/s)\n"
"p = sma × (1 − ecc²)         # 半通徑 (km)\n"
"dΩ/dt = −(3/2) × J2 × (RE/p)² × n × cos(inc)   # deg/s\n\n"
"draan_res = RAAN_raw_diff − (dΩ/dt × Δt_seconds)")
    doc.add_paragraph()

    add_heading(doc, "2.3 旗標閾值", level=2)
    add_table(doc,
        ["參數", "閾值", "說明"],
        [
            ["THR_DA_SM", "> 1.0 km",  "半長軸變化量"],
            ["THR_DI",    "> 0.02°",   "傾角變化量"],
            ["THR_DE",    "> 0.001",   "離心率變化量"],
            ["THR_DRAAN", "> 0.1°",    "J2 修正後 RAAN 殘差"],
        ],
        col_widths=[3.0, 3.0, 9.0])
    add_body(doc, "任一差分超過對應閾值，則該轉移被標記 flagged=True。")

    # ── 3. 衍生特徵 ──────────────────────────────────────────────────────────
    add_heading(doc, "3. 衍生特徵計算：classify_behavior() 前半段", level=1)

    add_heading(doc, "3.1 觸發維度（布林）", level=2)
    add_code(doc,
"has_da    = 任一 flagged 轉移的 |da| > 1.0 km\n"
"has_di    = 任一 flagged 轉移的 |di| > 0.02°\n"
"has_de    = 任一 flagged 轉移的 |de| > 0.001\n"
"has_draan = 任一 flagged 轉移的 |draan_res| > 0.1°")

    add_heading(doc, "3.2 幅度指標", level=2)
    add_code(doc,
"max_da  = flagged 轉移中 |da| 的最大值   （km）\n"
"net_da  = 所有轉移 da 的代數和           （21 天累積軌道漂移，km）")

    add_heading(doc, "3.3 單調性指標（識別大氣阻力 vs 主動機動）", level=2)
    add_body(doc,
        "大氣阻力造成的 sma 下降是持續不中斷的；主動機動通常有正有負、非單調。"
        "以最長連續負值序列長度作為判斷依據：")
    add_code(doc,
"neg_streak = 連續 da < −0.3 km 的最長序列長度\n"
"total_drop = abs(net_da) if net_da < 0\n\n"
"monotone_decay = (neg_streak ≥ 5)\n"
"             AND (total_drop > 5 km)\n"
"             AND (net_da < −3 km)")
    add_body(doc,
        "三個條件同時成立才判定為單調衰減，避免短暫拉力補償期被誤判。")

    add_heading(doc, "3.4 頻率指標", level=2)
    add_code(doc,
"flag_rate       = n_flagged / n_transitions    # 旗標轉移佔所有轉移的比例\n"
"burn_freq_per_day = n_flagged / 21             # 日均機動次數（觀測窗 21 天）")

    # ── 4. 決策樹 ────────────────────────────────────────────────────────────
    add_heading(doc, "4. 規則決策樹：classify_behavior() 後半段", level=1)
    add_body(doc,
        "分類器首先以推進系統標註作為最頂層分歧，再依據衍生特徵逐層判定。"
        "優先順序設計的關鍵：monotone_decay 最優先，確保大幅降軌的電推衛星"
        "（如 JILIN-01 GAOFEN 3J，net_da = −117 km）不被誤分為 OrbitLowering。")
    doc.add_paragraph()
    add_code(doc,
"propulsion_class ∈ {Electric_EP, Chemical, Micro/ColdGas, Hybrid/Other}?\n"
"│\n"
"├─ 否（無推進標註）\n"
"│   ├─ monotone_decay = True     → DragDecay          (純大氣阻力衰減 FP)\n"
"│   ├─ has_draan only            → UnknownRAANAnomaly  (RAAN 異常待查)\n"
"│   └─ 其他                      → UnknownFP           (未知 FP)\n"
"│\n"
"└─ 是（有推進標註）\n"
"    ├─ monotone_decay = True      → DragMakeup          (電推補償拖曳)\n"
"    ├─ net_da > +5 km             → OrbitRaising        (主動升軌)\n"
"    ├─ net_da < −5 km             → OrbitLowering       (主動降軌/離軌)\n"
"    ├─ has_da AND has_di          → ComplexManeuver     (複合機動)\n"
"    ├─ has_di only                → InclinationChange   (純傾角調整)\n"
"    ├─ has_draan only             → RAANMaintenance     (RAAN 維持)\n"
"    ├─ flag_rate > 25%            → Stationkeeping      (高頻小幅機動)\n"
"    ├─ n_flagged ≤ 3 AND max_da > 5 km → PhasingManeuver (少次大幅機動)\n"
"    ├─ flag_rate ≤ 15%            → DragMakeup          (低頻小幅補償)\n"
"    └─ 其他                       → Stationkeeping")
    doc.add_paragraph()

    add_heading(doc, "4.1 十種行為類別說明", level=2)
    add_table(doc,
        ["類別代號", "中文名稱", "核心判斷條件", "代表案例"],
        [
            ["DragMakeup",       "大氣阻力補償",
             "monotone_decay=T（有推進）\n或 flag_rate ≤ 15%",
             "758 顆 Starlink（主流）"],
            ["OrbitLowering",    "軌道降低/離軌",
             "net_da < −5 km（有推進）",
             "Starlink 離軌程序（252 顆）"],
            ["PhasingManeuver",  "相位調整",
             "n_flagged ≤ 3 且 max_da > 5 km",
             "Starlink 新部署（91 顆）"],
            ["OrbitRaising",     "軌道抬升",
             "net_da > +5 km（有推進）",
             "KUIPER-00262 +64 km（81 顆）"],
            ["Stationkeeping",   "站位保持",
             "flag_rate > 25%",
             "SAR/EO 精密星座（69 顆）"],
            ["UnknownFP",        "未知 FP",
             "無推進標註，非單調衰減",
             "OBJECT 系列（39 顆）"],
            ["DragDecay",        "大氣阻力衰減",
             "monotone_decay=T（無推進）",
             "RHOK-SAT 268 km（21 顆）"],
            ["InclinationChange","傾角調整",
             "has_di=T，has_da=F",
             "Starlink 面調整（6 顆）"],
            ["ComplexManeuver",  "複合機動",
             "has_da=T 且 has_di=T",
             "KUIPER-00262（4 顆）"],
            ["UnknownRAANAnomaly","RAAN 異常",
             "has_draan=T，無 da/di",
             "IXPE、ORS 5（2 顆）"],
        ],
        col_widths=[3.2, 3.2, 5.0, 4.5])

    # ── 5. 限制 ──────────────────────────────────────────────────────────────
    add_heading(doc, "5. 已知限制與潛在問題", level=1)

    add_heading(doc, "5.1 DragMakeup 定義過寬", level=2)
    add_body(doc,
        "monotone_decay=True 且有推進 → DragMakeup。"
        "但 JILIN-01 GAOFEN 3J（net_da = −117 km）是真正的終端降軌，"
        "語意上應為 OrbitLowering，目前被錯分為 DragMakeup。"
        "改進方向：加入 total_drop > 50 km 的門檻以區分「大幅終端降軌」與「持續阻力補償」。")

    add_heading(doc, "5.2 OrbitLowering 與 Starlink 離軌混淆", level=2)
    add_body(doc,
        "Starlink 離軌程序在 21 天窗口內執行 1–3 次大幅降軌，net_da 明顯為負，"
        "被分到 OrbitLowering。但部分短期相位調整後尚未補回，也落在同一類別，"
        "語意不完全相同但特徵無法區分。")

    add_heading(doc, "5.3 PhasingManeuver 的 n_flagged ≤ 3 是硬編碼閾值", level=2)
    add_body(doc,
        "Starlink 新星座部署通常 1–2 次燒到位；CSS 空間站軌道維持燒同樣可能只有 1 次但幅度大，"
        "兩者特徵相似但語意不同，目前無法區分。")

    add_heading(doc, "5.4 monotone_decay 未考慮 B* 值", level=2)
    add_body(doc,
        "原始演算法改進記憶（project_algorithm_improvements.md）提到應同時參考 B*（氣動阻力係數），"
        "目前版本只用 da 序列本身。對 B* 大的高阻力小衛星判斷較準，"
        "對中等高度的標準衛星（B* ≈ 正常）可能有誤判。")

    add_heading(doc, "5.5 觀測窗口短（21 天）造成 Recall 低", level=2)
    add_body(doc,
        "本次驗證 Recall ≈ 12.4%，主要原因如下：")
    items = [
        "觀測窗口短（21 天）：電推衛星每次燃燒僅 1–3 km，21 天內不一定執行機動",
        "閾值偏高：THR_DA_SM = 1.0 km 過濾掉微小機動，低軌 Starlink 電推量約 0.1–0.5 km/burn",
        "TLE 更新頻率不足：機動發生於兩個 TLE epoch 之間無法偵測",
        "多機動抵消：升軌後降軌的累積 da 趨近零，被視為無機動",
    ]
    for item in items:
        p = doc.add_paragraph(item, style="List Bullet")
        p.runs[0].font.size = Pt(11)

    # ── 6. 輸出格式 ──────────────────────────────────────────────────────────
    add_heading(doc, "6. 輸出格式", level=1)

    add_heading(doc, "6.1 maneuver_behavior_classified.csv（每列一顆衛星）", level=2)
    add_table(doc,
        ["欄位", "型別", "說明"],
        [
            ["norad_id",          "str",   "NORAD 衛星編號"],
            ["sat_name",          "str",   "衛星名稱"],
            ["propulsion_class",  "str",   "推進類別（Electric_EP / Chemical / ...）"],
            ["behavior",          "str",   "行為類別（10 種）"],
            ["net_da_km",         "float", "21 天累積 Δa（km），正值升軌負值降軌"],
            ["max_da_km",         "float", "單次轉移最大 |Δa|（km）"],
            ["max_di_deg",        "float", "單次轉移最大 |Δi|（deg）"],
            ["n_flagged",         "int",   "觀測期間旗標轉移次數"],
            ["n_transitions",     "int",   "觀測期間總轉移次數"],
            ["flag_rate",         "float", "旗標率 = n_flagged / n_transitions"],
            ["burn_freq_per_day", "float", "日均機動次數 = n_flagged / 21"],
            ["monotone_decay",    "bool",  "是否判定為單調衰減（大氣阻力）"],
            ["triggers",          "str",   "觸發維度組合（da / di / de / draan）"],
            ["mean_sma_km",       "float", "平均 SMA（km），換算高度需 −6378.137 km）"],
            ["sub_tags",          "str",   "細分標籤（da_max、net、monotone_drop 等）"],
        ],
        col_widths=[4.2, 1.8, 9.0])

    add_heading(doc, "6.2 maneuver_behavior_report.md（Markdown 報告）", level=2)
    add_table(doc,
        ["章節", "內容"],
        [
            ["§1 行為類別總覽",        "10 種行為的顆數與佔比"],
            ["§2 推進類別 × 行為分佈", "交叉表（propulsion_class × behavior）"],
            ["§3 各類別詳細說明",       "每類的統計摘要 + 前 15 顆代表衛星"],
            ["§4 特殊案例分析",         "da 最大 Top5、旗標最多 Top5、高頻燃燒 Top10"],
            ["§5 殘餘 FP 分析",         "62 顆無推進但被偵測衛星的完整列表"],
            ["§6 訓練資料建議",         "各類行為是否適合納入訓練正例，以及 Recall 偏低原因"],
        ],
        col_widths=[4.5, 11.0])

    # ── 7. 呼叫鏈 ────────────────────────────────────────────────────────────
    add_heading(doc, "7. 完整呼叫鏈", level=1)
    add_code(doc,
"main()\n"
" ├─ pd.read_csv(validation_full.csv)   → 過濾 maneuver=True → 1323 顆 NORAD\n"
" ├─ pd.read_csv(annotations_leo_full.csv) → 建立 norad_id 索引\n"
" ├─ duckdb.connect(space_db.duckdb)\n"
" │   └─ SELECT sma_km, ecc, inc, raan FROM tle_table\n"
" │      WHERE norad IN (...) AND date BETWEEN 2026-05-01 AND 2026-05-22\n"
" │      → 62,005 筆 TLE 記錄\n"
" └─ for norad in norads:\n"
"     ├─ detect_transitions(sat_tle)\n"
"     │   └─ for i in range(1, len(tle)):\n"
"     │       ├─ da  = sma[i] − sma[i-1]\n"
"     │       ├─ di  = inc[i] − inc[i-1]\n"
"     │       ├─ de  = ecc[i] − ecc[i-1]\n"
"     │       ├─ draan_res = angle_diff(raan[i], raan[i-1]) − j2_raan_rate() × dt\n"
"     │       └─ flagged = any threshold exceeded\n"
"     ├─ classify_behavior(tr, {propulsion_class})\n"
"     │   ├─ 計算 has_da, has_di, has_de, has_draan, net_da, max_da\n"
"     │   ├─ 計算 neg_streak → monotone_decay\n"
"     │   ├─ 計算 flag_rate, burn_freq\n"
"     │   └─ 決策樹 → behavior 類別\n"
"     └─ records.append(result)\n"
" ├─ DataFrame(records).to_csv(maneuver_behavior_classified.csv)\n"
" └─ generate_report(classified)\n"
"     └─ Path(OUT_MD).write_text(report, encoding='utf-8')")

    # ── 儲存 ─────────────────────────────────────────────────────────────────
    doc.save(OUT_PATH)
    print(f"[DONE] {OUT_PATH}")


if __name__ == "__main__":
    build_doc()
