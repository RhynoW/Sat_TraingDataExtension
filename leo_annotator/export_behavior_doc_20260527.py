"""
export_behavior_doc_20260527.py
================================
輸出 20260527 版架構說明 Word 文件，包含：
  - validate_annotations.py 時間範圍確認
  - 混淆矩陣原理（由淺入深）
  - TLE 偵測閾值設定依據
  - analyze_maneuver_behavior.py 行為分類特徵條件
  - 決策樹優先順序設計
  - 三大邊界難題
  - 分析結果摘要
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_PATH = r"output/analyze_maneuver_behavior_py_20260527架構說明.docx"


# ── 格式輔助 ─────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def add_heading(doc, text, level=1):
    p = doc.add_heading(text, level=level)
    color = RGBColor(0x1F, 0x49, 0x7D) if level == 1 else RGBColor(0x2E, 0x74, 0xB5)
    if p.runs:
        p.runs[0].font.color.rgb = color
    return p


def add_body(doc, text):
    p = doc.add_paragraph(text)
    if p.runs:
        p.runs[0].font.size = Pt(11)
    return p


def add_bullet(doc, text):
    p = doc.add_paragraph(text, style="List Bullet")
    if p.runs:
        p.runs[0].font.size = Pt(11)
    return p


def add_code(doc, text):
    p = doc.add_paragraph()
    p.paragraph_format.left_indent = Cm(1)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F2F2F2")
    pPr.append(shd)
    return p


def add_table(doc, headers, rows, col_widths=None, header_color="2E74B5"):
    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = h
        if cell.paragraphs[0].runs:
            run = cell.paragraphs[0].runs[0]
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.size = Pt(10)
        set_cell_bg(cell, header_color)
    for ri, row in enumerate(rows):
        tr = table.rows[ri + 1]
        bg = "F2F2F2" if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row):
            cell = tr.cells[ci]
            cell.text = str(val)
            if cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].font.size = Pt(10)
            if bg == "F2F2F2":
                set_cell_bg(cell, bg)
    if col_widths:
        for col_i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[col_i].width = Cm(w)
    return table


# ── 主文件 ───────────────────────────────────────────────────────────────────

def build_doc():
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Cm(2.0)
        section.bottom_margin = Cm(2.0)
        section.left_margin   = Cm(2.5)
        section.right_margin  = Cm(2.5)

    # ── 封面 ─────────────────────────────────────────────────────────────────
    doc.add_paragraph()
    title = doc.add_paragraph()
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run = title.add_run("LEO 衛星軌道機動偵測與行為分類")
    run.bold = True
    run.font.size = Pt(20)
    run.font.color.rgb = RGBColor(0x1F, 0x49, 0x7D)

    sub = doc.add_paragraph()
    sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run2 = sub.add_run("混淆矩陣原理 × 行為分類特徵條件 架構說明")
    run2.font.size = Pt(15)
    run2.font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)

    meta = doc.add_paragraph()
    meta.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run3 = meta.add_run("版本：2026-05-27　　分析資料期間：2026-05-01 ～ 2026-05-27（26 天）")
    run3.font.size = Pt(10)
    run3.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    meta2 = doc.add_paragraph()
    meta2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    run4 = meta2.add_run("涵蓋腳本：validate_annotations.py、analyze_maneuver_behavior.py")
    run4.font.size = Pt(10)
    run4.font.color.rgb = RGBColor(0x80, 0x80, 0x80)

    doc.add_paragraph()
    p_line = doc.add_paragraph("─" * 72)
    if p_line.runs:
        p_line.runs[0].font.color.rgb = RGBColor(0x2E, 0x74, 0xB5)
    doc.add_paragraph()

    # ── 0. 文件目的 ──────────────────────────────────────────────────────────
    add_heading(doc, "0. 文件目的", level=1)
    add_body(doc,
        "本文件說明 LEO 衛星軌道機動偵測管線中兩個核心腳本的運作邏輯：\n"
        "①  validate_annotations.py — 以 TLE 差分偵測機動並計算混淆矩陣，驗證推進系統標註品質；\n"
        "②  analyze_maneuver_behavior.py — 對偵測到機動的衛星進行行為分類，輸出 10 種行為類別。\n"
        "重點聚焦在：混淆矩陣的定義原理、偵測閾值的設定依據、行為分類特徵條件的量化邏輯，\n"
        "以及三大分類邊界難題的設計決策。")

    # ── 1. 分析時間範圍 ──────────────────────────────────────────────────────
    add_heading(doc, "1. 分析時間範圍確認（validate_annotations.py）", level=1)
    add_code(doc,
"# validate_annotations.py 第 23–24 行\n"
"DATE_START = \"2026-05-01\"\n"
"DATE_END   = \"2026-05-27\"   # 觀測窗口：26 天")
    add_body(doc,
        "觀測窗口為 2026-05-01 至 2026-05-27，共 26 天。\n"
        "查詢條件為 DuckDB tle_table 的 date_tag BETWEEN '2026-05-01' AND '2026-05-27 23:59:59'。\n"
        "TLE 對間隔超過 7 天的轉移會被跳過，以避免資料缺口造成誤判。")
    add_table(doc,
        ["參數", "值", "說明"],
        [
            ["DATE_START",    "2026-05-01",   "觀測期起始日"],
            ["DATE_END",      "2026-05-27",   "觀測期結束日（原 2026-05-22，已延長至 27 日）"],
            ["觀測長度",       "26 天",         "用於計算 burn_freq_per_day"],
            ["最長 TLE 間隔",  "7 天",          "超過則跳過，避免缺測誤判"],
            ["標註總數",       "14,090 顆",      "annotations_leo_full.csv 全量"],
            ["有效分析",       "約 13,700+ 顆",  "TLE ≥ 3 筆且有轉移記錄"],
        ],
        col_widths=[3.5, 3.5, 9.0])

    # ── 2. 混淆矩陣原理 ──────────────────────────────────────────────────────
    add_heading(doc, "2. 混淆矩陣原理（由淺入深）", level=1)

    add_heading(doc, "2.1 一句話概念", level=2)
    add_body(doc,
        "混淆矩陣比較「我們給衛星貼的標籤」與「TLE 資料顯示的行為」是否一致。\n"
        "標籤（Ground Truth）= 衛星是否有推進系統；預測（Prediction）= TLE 差分是否偵測到軌道改變。")

    add_heading(doc, "2.2 Ground Truth：推進系統標註", level=2)
    add_body(doc,
        "Ground Truth 直接來自 annotations_leo_full.csv 的 propulsion_class 欄位。\n"
        "凡 propulsion_class ∈ {Electric_EP, Chemical, Micro/ColdGas, Hybrid/Other} 者，\n"
        "視為「有推進能力」（has_prop = True）；propulsion_class = 'None' 者視為「無推進」。")
    add_code(doc,
"# validate_annotations.py 第 241–244 行\n"
"has_prop = ok[\"propulsion_class\"].isin(\n"
"    [\"Electric_EP\", \"Chemical\", \"Micro/ColdGas\", \"Hybrid/Other\"]) & \\\n"
"    ok[\"propulsion_class\"].notna()\n"
"\n"
"# 有推進 = True  →  Ground Truth = Positive\n"
"# 無推進 = False →  Ground Truth = Negative")

    add_heading(doc, "2.3 Prediction：TLE 機動偵測", level=2)
    add_body(doc,
        "Prediction 來自 detect_maneuvers_for_sat() 的輸出：\n"
        "對衛星在觀測期間的所有相鄰 TLE，計算軌道差分，任一差分超過閾值即標記 flagged=True；\n"
        "若該衛星有至少 1 個 flagged 轉移，則 maneuver_detected = True（預測 = 發生機動）。")
    add_code(doc,
"# validate_annotations.py 第 189–191 行\n"
"flagged             = trans[trans[\"flagged\"]]\n"
"rec[\"n_flagged\"]         = int(len(flagged))\n"
"rec[\"maneuver_detected\"] = bool(len(flagged) > 0)\n"
"\n"
"# maneuver_detected = True  →  Prediction = Positive\n"
"# maneuver_detected = False →  Prediction = Negative")

    add_heading(doc, "2.4 四格矩陣定義", level=2)
    add_table(doc,
        ["縮寫", "全名", "意義", "本次數量"],
        [
            ["TP", "True Positive",  "有推進 + 偵到機動（正確陽性）",    "~154 顆"],
            ["FN", "False Negative", "有推進 + 未偵到機動（漏報）",      "~1,089 顆"],
            ["FP", "False Positive", "無推進 + 偵到機動（誤報）",        "62 顆"],
            ["TN", "True Negative",  "無推進 + 未偵到機動（正確陰性）",   "~12,000+ 顆"],
        ],
        col_widths=[1.8, 3.5, 6.5, 3.0])
    add_code(doc,
"# validate_annotations.py 第 245–248 行\n"
"tp = int(ok[ has_prop  &  ok[\"maneuver_detected\"]].shape[0])\n"
"fn = int(ok[ has_prop  & ~ok[\"maneuver_detected\"]].shape[0])\n"
"fp = int(ok[~has_prop  &  ok[\"maneuver_detected\"]].shape[0])\n"
"tn = int(ok[~has_prop  & ~ok[\"maneuver_detected\"]].shape[0])\n"
"\n"
"precision = tp / (tp + fp)   # 偵測到的機動中，真的有推進的比例\n"
"recall    = tp / (tp + fn)   # 所有有推進衛星中，被成功偵測的比例\n"
"f1        = 2 * precision * recall / (precision + recall)")

    add_heading(doc, "2.5 Precision = 95.3%、Recall = 12.4% 的解讀", level=2)
    add_body(doc,
        "Precision 95.3% 代表：TLE 偵測到的機動，有 95.3% 確實來自有推進的衛星，誤報率極低。\n"
        "Recall 12.4% 代表：有推進的衛星中，只有 12.4% 在 26 天內被 TLE 偵測到機動。\n"
        "Recall 偏低是結構性問題，不代表腳本有 bug，原因如下：")
    add_bullet(doc, "「有推進系統」≠「這 26 天內恰好在執行機動」——大多數電推衛星每次燃燒僅 0.1–0.3 km，"
               "間隔可能超過觀測窗口")
    add_bullet(doc, "閾值 THR_DA_SM = 1.0 km 約等於 TLE 測量噪聲的 2× ，低於此的機動無法與噪聲區分")
    add_bullet(doc, "TLE 更新頻率（通常 1–2 次/天）不足以捕捉兩次 TLE 之間發生的燃燒")
    add_bullet(doc, "相互抵消的機動（先升後降）累積 Δa 趨近 0，被誤判為無機動")

    # ── 3. 偵測閾值設定依據 ──────────────────────────────────────────────────
    add_heading(doc, "3. TLE 偵測閾值設定依據", level=1)
    add_body(doc,
        "偵測閾值設定的核心原則：高於 TLE 測量噪聲的 2 倍，低於典型機動幅度。\n"
        "TLE 本身是從雷達觀測擬合出的軌道元素，存在測量誤差；"
        "若閾值太低，噪聲會被誤判為機動（高 FP）；若太高，真正的機動會被漏掉（高 FN）。")
    add_table(doc,
        ["參數", "閾值", "TLE 噪聲估計", "設定依據"],
        [
            ["THR_DA_SM",  "> 1.0 km",   "0.3–0.5 km",  "約 2× 噪聲；典型 LEO 電推單次燃燒 0.1–3 km"],
            ["THR_DI",     "> 0.02°",    "0.005–0.01°", "約 2× 噪聲；傾角調整通常 0.1–1°"],
            ["THR_DE",     "> 0.001",    "0.0003–0.0005","約 2× 噪聲；主動機動通常 0.005–0.05"],
            ["THR_DRAAN",  "> 0.1°",     "0.03–0.05°",  "J2 校正後殘差；主動 RAAN 調整通常 > 0.3°"],
        ],
        col_widths=[2.8, 2.2, 3.5, 7.0])
    add_body(doc,
        "J2 RAAN 修正（j2_raan_rate_deg_s）扣除地球扁率引起的自然進動，"
        "使 RAAN 閾值真正代表主動側噴推力，而非地球重力場的被動效應。")
    add_code(doc,
"# validate_annotations.py 第 41–45 行\n"
"def j2_raan_rate_deg_s(sma_km, ecc, inc_deg):\n"
"    n = sqrt(MU / sma_km**3)          # 平均角速度 (rad/s)\n"
"    p = sma_km * (1 - ecc**2)         # 半通徑 (km)\n"
"    return degrees(-1.5 * J2 * (RE/p)**2 * n * cos(radians(inc_deg)))\n"
"\n"
"draan_res = RAAN_raw_diff - j2_raan_rate_deg_s(...) * dt_seconds")

    # ── 4. 行為分類特徵條件 ──────────────────────────────────────────────────
    add_heading(doc, "4. 行為分類特徵條件（analyze_maneuver_behavior.py）", level=1)
    add_body(doc,
        "行為分類器的輸入特徵由 detect_transitions() 算出的轉移資料衍生而來，"
        "分為四類：觸發維度、幅度指標、單調性指標、頻率指標。")

    add_heading(doc, "4.1 觸發維度（布林，判斷機動發生在哪個軌道元素）", level=2)
    add_code(doc,
"has_da    = 任一 flagged 轉移的 |da|        > 1.0 km\n"
"has_di    = 任一 flagged 轉移的 |di|        > 0.02°\n"
"has_de    = 任一 flagged 轉移的 |de|        > 0.001\n"
"has_draan = 任一 flagged 轉移的 |draan_res| > 0.1°")
    add_body(doc,
        "這四個布林值告訴分類器：衛星是在調整「軌道高度、軌道面傾角、橢圓形狀、升交點方向」"
        "中的哪一個（或哪些組合）。")

    add_heading(doc, "4.2 幅度指標（量化機動規模）", level=2)
    add_code(doc,
"max_da  = flagged 轉移中 |da| 的最大值  （km）  — 單次最大機動幅度\n"
"net_da  = 所有轉移 da 的代數和           （km）  — 26 天累積軌道高度變化\n"
"          正值 → 衛星整體在升軌\n"
"          負值 → 衛星整體在降軌")
    add_body(doc,
        "net_da ±5 km 是區分「主動換軌」與「小幅調整」的關鍵門檻：\n"
        "OrbitRaising: net_da > +5 km（衛星在這 26 天明顯爬升）\n"
        "OrbitLowering: net_da < -5 km（衛星在這 26 天明顯下降）\n"
        "其餘的「小幅機動」則進一步依頻率和次數分類。")

    add_heading(doc, "4.3 單調性指標（識別大氣阻力 vs 主動機動）", level=2)
    add_body(doc,
        "大氣阻力造成的 SMA 下降是「持續不中斷的單向衰減」；\n"
        "主動機動通常是「間歇性的，有時向上有時向下，非單調」。\n"
        "以最長連續負 da 序列長度捕捉這個物理差異：")
    add_code(doc,
"neg_streak  = 所有轉移中 da < -0.3 km 連續出現的最長序列長度\n"
"total_drop  = |net_da| （僅在 net_da < 0 時有意義）\n"
"\n"
"monotone_decay = True  當且僅當：\n"
"    neg_streak ≥ 5       # 連續 5 次以上持續下降（排除偶然噪聲）\n"
"  AND total_drop > 5 km  # 累積下降超過 5 km（排除短暫波動）\n"
"  AND net_da < -3 km     # 整體確實在下降（排除升降抵消的情況）")
    add_body(doc,
        "三個條件必須同時成立，主要是為了排除「短暫執行機動後恢復」的情況："
        "若衛星先連續降軌 5 次（neg_streak=5）但最後又補回來，"
        "net_da 會接近 0，不會被誤判為大氣衰減。")

    add_heading(doc, "4.4 頻率指標（判斷機動模式：持續站位 vs 偶發調整）", level=2)
    add_code(doc,
"flag_rate       = n_flagged / n_transitions     # 旗標轉移占所有轉移的比例\n"
"burn_freq_per_day = n_flagged / 26              # 日均機動次數（26 天觀測窗）\n"
"\n"
"Stationkeeping  ← flag_rate > 25%   # 每 4 次 TLE 更新就有 1 次機動：高頻持續調整\n"
"DragMakeup      ← flag_rate ≤ 15%   # 偶爾補一下：低頻小幅電推補償")

    # ── 5. 決策樹優先順序 ────────────────────────────────────────────────────
    add_heading(doc, "5. 決策樹優先順序設計邏輯", level=1)
    add_body(doc,
        "分類器採用規則決策樹（rule-based decision tree），"
        "特徵按優先順序依序判斷。優先順序的設計關鍵如下：")
    add_code(doc,
"propulsion_class ∈ {Electric_EP, Chemical, Micro/ColdGas, Hybrid/Other}?\n"
"│\n"
"├─ 否（無推進標註）\n"
"│   ├─ [優先1] monotone_decay = True     → DragDecay          (純大氣阻力衰減)\n"
"│   ├─ [優先2] has_draan=T, has_da=F     → UnknownRAANAnomaly (待查)\n"
"│   └─ [其他]                            → UnknownFP          (未知 FP)\n"
"│\n"
"└─ 是（有推進標註）\n"
"    ├─ [優先1] monotone_decay = True      → DragMakeup         (電推補償拖曳)\n"
"    ├─ [優先2] net_da > +5 km             → OrbitRaising       (主動升軌)\n"
"    ├─ [優先3] net_da < −5 km             → OrbitLowering      (主動降軌)\n"
"    ├─ [優先4] has_da AND has_di          → ComplexManeuver    (複合機動)\n"
"    ├─ [優先5] has_di only               → InclinationChange   (純傾角調整)\n"
"    ├─ [優先6] has_draan only            → RAANMaintenance     (RAAN 維持)\n"
"    ├─ [優先7] flag_rate > 25%            → Stationkeeping     (高頻小幅)\n"
"    ├─ [優先8] n_flagged ≤ 3 AND max_da > 5 km → PhasingManeuver\n"
"    ├─ [優先9] flag_rate ≤ 15%            → DragMakeup         (低頻補償)\n"
"    └─ [其他]                             → Stationkeeping")

    add_heading(doc, "5.1 為什麼 monotone_decay 要排在最優先？", level=2)
    add_body(doc,
        "設計動機來自 JILIN-01 GAOFEN 3J（NORAD 58039）這個真實案例：\n"
        "net_da = -117 km（整個觀測期間持續大幅降軌），\n"
        "如果 monotone_decay 不先判斷，會掉進優先3（net_da < -5 km → OrbitLowering），\n"
        "但這顆衛星實際是在執行大氣阻力補償的持續電推燃燒，不是主動離軌。\n"
        "monotone_decay 放在最前面，確保「持續衰減模式」比「淨位移方向」更優先。")

    add_heading(doc, "5.2 PhasingManeuver 的 n_flagged ≤ 3 條件", level=2)
    add_body(doc,
        "Starlink 新部署時通常只燒 1–2 次就到位（n_flagged 很少），"
        "但每次幅度很大（max_da 可達 10–30 km）。\n"
        "這個組合（少次 + 大幅）在物理上代表「一次性相位調整」，"
        "與「高頻持續小幅的站位保持」明顯不同。\n"
        "n_flagged ≤ 3 是從 91 顆分類結果反推出的經驗門檻。")

    # ── 6. 三大邊界難題 ──────────────────────────────────────────────────────
    add_heading(doc, "6. 三大分類邊界難題", level=1)

    add_heading(doc, "難題 1：DragMakeup vs OrbitLowering", level=2)
    add_body(doc,
        "問題：都是 net_da < 0，都有 has_da = True。\n"
        "區分依據：大氣阻力補償的 SMA 下降是「單調連續的」；主動降軌通常是「間歇性的」，\n"
        "下降後可能停止或有小幅補回。\n"
        "monotone_decay=True 代表連續 ≥ 5 個 TLE pair 都是持續下降，是大氣補償特徵；\n"
        "反之，net_da < -5 km 且 monotone_decay=False，則是間歇性主動降軌。\n"
        "殘餘問題：JILIN-01 GAOFEN 3J 被歸入 DragMakeup（語意應為終端降軌），"
        "可加入 total_drop > 50 km 進一步區分。")

    add_heading(doc, "難題 2：PhasingManeuver vs 單次軌道維持燃燒", level=2)
    add_body(doc,
        "問題：兩者都是 n_flagged ≤ 3 且 max_da > 5 km。\n"
        "前者（相位調整）：Starlink 新星座部署，燒 1–2 次到達目標相位；\n"
        "後者（軌道維持燒）：CSS 空間站 / 某些 GEO 衛星每隔很長時間才燒一次，幅度不小。\n"
        "目前無法區分，因為兩者的 TLE 差分特徵相同。\n"
        "改進方向：引入衛星類型標籤（LEO 小衛星 vs ISS/CSS 空間站）作為額外篩選條件。")

    add_heading(doc, "難題 3：DragDecay（無推進衰減）vs 電推耗盡後的衰減", level=2)
    add_body(doc,
        "問題：DragDecay 要求 propulsion_class = 'None' + monotone_decay = True；\n"
        "但部分衛星電推耗盡後，行為與無推進衛星完全相同，且標註仍為 Electric_EP。\n"
        "如果標註正確（電推耗盡仍標 Electric_EP），這類衛星會被歸入 DragMakeup；\n"
        "如果標註未能及時更新，則仍被歸入 DragMakeup 但語意錯誤。\n"
        "根本解法：需要時序上的發射年份或電推壽命資訊，目前暫以人工核查補充。")

    # ── 7. 分析結果摘要 ──────────────────────────────────────────────────────
    add_heading(doc, "7. 分析結果摘要（截至 2026-05-27）", level=1)

    add_heading(doc, "7.1 驗證指標", level=2)
    add_table(doc,
        ["指標", "值", "說明"],
        [
            ["Precision", "95.3%", "偵測到的機動中確實有推進系統的比例"],
            ["Recall",    "12.4%", "有推進衛星中在 26 天內被偵測到的比例"],
            ["F1-score",  "21.9%", "Precision 與 Recall 的調和平均"],
            ["TP",        "~154 顆", "有推進 + 偵到機動"],
            ["FN",        "~1,089 顆", "有推進 + 未偵到（觀測窗口內未機動）"],
            ["FP",        "62 顆",  "無推進 + 偵到機動（人工核查後剩餘）"],
            ["TN",        "~12,000+ 顆", "無推進 + 未偵到"],
        ],
        col_widths=[3.0, 3.0, 10.0])

    add_heading(doc, "7.2 行為分類結果（共 1,323 顆偵測到機動）", level=2)
    add_table(doc,
        ["行為類別", "顆數", "占比", "說明"],
        [
            ["DragMakeup",        "758", "57.3%", "電推補償大氣阻力（Starlink 主流模式）"],
            ["OrbitLowering",     "252", "19.0%", "主動降軌或離軌程序"],
            ["PhasingManeuver",    "91",  "6.9%", "新部署相位調整（少次大幅燃燒）"],
            ["OrbitRaising",       "81",  "6.1%", "主動升軌（Kuiper 等新部署）"],
            ["Stationkeeping",     "69",  "5.2%", "高頻小幅持續站位保持"],
            ["UnknownFP",          "39",  "2.9%", "無推進標註但有機動訊號（殘餘 FP）"],
            ["DragDecay",          "21",  "1.6%", "純大氣阻力衰減（極低軌無推進）"],
            ["InclinationChange",   "6",  "0.5%", "純傾角調整"],
            ["ComplexManeuver",     "4",  "0.3%", "同時有 da 和 di 的複合機動"],
            ["UnknownRAANAnomaly",  "2",  "0.2%", "RAAN 異常、da/di 正常（待查）"],
        ],
        col_widths=[4.0, 2.0, 2.0, 8.0])

    add_heading(doc, "7.3 輸出檔案", level=2)
    add_table(doc,
        ["檔案", "說明"],
        [
            ["output/validation_full.csv",
             "14,090 顆衛星逐顆驗證結果（tle_status、maneuver_detected、n_flagged 等）"],
            ["output/validation_report_full.txt",
             "驗證摘要報告（混淆矩陣、分類別偵測率）"],
            ["output/maneuver_behavior_classified.csv",
             "1,323 顆機動衛星的行為分類結果（behavior、net_da、monotone_decay 等 15 欄）"],
            ["output/maneuver_behavior_report.md",
             "Markdown 格式分析報告（6 章節含交叉表、特殊案例、FP 分析）"],
        ],
        col_widths=[6.5, 9.5])

    # ── 儲存 ─────────────────────────────────────────────────────────────────
    doc.save(OUT_PATH)
    print(f"[DONE] {OUT_PATH}")


if __name__ == "__main__":
    build_doc()
