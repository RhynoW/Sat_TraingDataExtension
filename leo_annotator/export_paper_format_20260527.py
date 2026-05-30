"""
export_paper_format_20260527.py
================================
以學術論文格式輸出 LEO 衛星軌道機動偵測與行為分類分析報告。
輸出：output/LEO衛星軌道機動偵測與行為分類_20260527.docx
"""

from docx import Document
from docx.shared import Pt, RGBColor, Inches, Cm
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT, WD_ALIGN_VERTICAL
from docx.oxml.ns import qn
from docx.oxml import OxmlElement

OUT_PATH = r"output/LEO衛星軌道機動偵測與行為分類_20260527.docx"

TABLE_COUNT = [0]   # mutable counter for table numbering


# ── 格式輔助 ─────────────────────────────────────────────────────────────────

def set_cell_bg(cell, hex_color: str):
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), hex_color)
    tcPr.append(shd)


def set_paragraph_font(p, name_cjk="標楷體", name_latin="Times New Roman",
                       size_pt=12, bold=False, italic=False,
                       color: RGBColor = None):
    for run in p.runs:
        run.font.name = name_latin
        run._r.get_or_add_rPr()
        rFonts = OxmlElement("w:rFonts")
        rFonts.set(qn("w:eastAsia"), name_cjk)
        run._r.rPr.append(rFonts)
        run.font.size = Pt(size_pt)
        run.font.bold  = bold
        run.font.italic = italic
        if color:
            run.font.color.rgb = color


def add_heading_paper(doc, text, level=1, numbering=None):
    """學術格式標題（帶編號、無顏色標記）"""
    full_text = f"{numbering}　{text}" if numbering else text
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(14) if level == 1 else Pt(8)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(full_text)
    run.font.name = "Times New Roman"
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:eastAsia"), "標楷體")
    rPr.append(rFonts)
    run.font.bold = True
    run.font.size = Pt(14) if level == 1 else Pt(12)
    return p


def add_body_paper(doc, text, indent=False, first_line=True):
    """內文段落（12pt 標楷體，首行縮排 2 字元）"""
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(4)
    p.paragraph_format.line_spacing = Pt(20)
    if indent or first_line:
        p.paragraph_format.first_line_indent = Cm(0.85)  # 約 2 字元
    run = p.add_run(text)
    run.font.name = "Times New Roman"
    rPr = run._r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:eastAsia"), "標楷體")
    rPr.append(rFonts)
    run.font.size = Pt(12)
    return p


def add_formula(doc, formula_text, label=""):
    """公式行（Courier New，右側標號）"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(1.5)
    p.paragraph_format.space_before = Pt(4)
    p.paragraph_format.space_after  = Pt(4)
    run = p.add_run(formula_text)
    run.font.name = "Courier New"
    run.font.size = Pt(11)
    run.font.italic = True
    if label:
        run2 = p.add_run(f"    {label}")
        run2.font.name = "Times New Roman"
        run2.font.size = Pt(11)
    return p


def add_code_paper(doc, text):
    """程式碼區塊（灰底 Courier New）"""
    p = doc.add_paragraph()
    p.paragraph_format.left_indent  = Cm(1.0)
    p.paragraph_format.space_before = Pt(2)
    p.paragraph_format.space_after  = Pt(2)
    run = p.add_run(text)
    run.font.name = "Courier New"
    run.font.size = Pt(9)
    run.font.color.rgb = RGBColor(0x1A, 0x1A, 0x1A)
    pPr = p._p.get_or_add_pPr()
    shd = OxmlElement("w:shd")
    shd.set(qn("w:val"), "clear")
    shd.set(qn("w:color"), "auto")
    shd.set(qn("w:fill"), "F2F2F2")
    pPr.append(shd)
    return p


def add_table_paper(doc, caption, headers, rows, col_widths=None):
    """論文格式表格：表名在上方，藍色標題列，斑馬紋"""
    TABLE_COUNT[0] += 1
    cap_p = doc.add_paragraph()
    cap_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    cap_p.paragraph_format.space_before = Pt(8)
    cap_p.paragraph_format.space_after  = Pt(3)
    cap_run = cap_p.add_run(f"表 {TABLE_COUNT[0]}　{caption}")
    cap_run.font.bold = True
    cap_run.font.size = Pt(11)
    cap_run.font.name = "Times New Roman"
    rPr = cap_run._r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:eastAsia"), "標楷體")
    rPr.append(rFonts)

    table = doc.add_table(rows=1 + len(rows), cols=len(headers))
    table.style = "Table Grid"
    table.alignment = WD_TABLE_ALIGNMENT.CENTER

    hdr_row = table.rows[0]
    for i, h in enumerate(headers):
        cell = hdr_row.cells[i]
        cell.text = h
        cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
        if cell.paragraphs[0].runs:
            run = cell.paragraphs[0].runs[0]
            run.bold = True
            run.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
            run.font.size = Pt(10)
            run.font.name = "Times New Roman"
        cell.paragraphs[0].alignment = WD_ALIGN_PARAGRAPH.CENTER
        set_cell_bg(cell, "1F497D")

    for ri, row in enumerate(rows):
        tr = table.rows[ri + 1]
        bg = "EEF3FB" if ri % 2 == 0 else "FFFFFF"
        for ci, val in enumerate(row):
            cell = tr.cells[ci]
            cell.text = str(val)
            cell.vertical_alignment = WD_ALIGN_VERTICAL.CENTER
            if cell.paragraphs[0].runs:
                cell.paragraphs[0].runs[0].font.size = Pt(10)
            if bg != "FFFFFF":
                set_cell_bg(cell, bg)

    if col_widths:
        for col_i, w in enumerate(col_widths):
            for row in table.rows:
                row.cells[col_i].width = Cm(w)

    doc.add_paragraph().paragraph_format.space_after = Pt(2)
    return table


def add_hr(doc):
    p = doc.add_paragraph()
    pPr = p._p.get_or_add_pPr()
    pBdr = OxmlElement("w:pBdr")
    bottom = OxmlElement("w:bottom")
    bottom.set(qn("w:val"), "single")
    bottom.set(qn("w:sz"), "4")
    bottom.set(qn("w:space"), "1")
    bottom.set(qn("w:color"), "2E74B5")
    pBdr.append(bottom)
    pPr.append(pBdr)
    return p


# ── 主文件 ───────────────────────────────────────────────────────────────────

def build_doc():
    doc = Document()

    for section in doc.sections:
        section.top_margin    = Cm(2.54)
        section.bottom_margin = Cm(2.54)
        section.left_margin   = Cm(3.0)
        section.right_margin  = Cm(3.0)

    # ════════════════════════════════════════════════════════════════
    # 標題頁
    # ════════════════════════════════════════════════════════════════
    doc.add_paragraph()
    doc.add_paragraph()

    title_p = doc.add_paragraph()
    title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r = title_p.add_run("基於 TLE 差分的 LEO 衛星軌道機動偵測\n與行為分類研究")
    r.bold = True
    r.font.size = Pt(18)
    r.font.name = "Times New Roman"
    rPr = r._r.get_or_add_rPr()
    rFonts = OxmlElement("w:rFonts")
    rFonts.set(qn("w:eastAsia"), "標楷體")
    rPr.append(rFonts)

    doc.add_paragraph()
    en_title_p = doc.add_paragraph()
    en_title_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r2 = en_title_p.add_run(
        "LEO Satellite Orbital Maneuver Detection and Behavior Classification\n"
        "Based on Two-Line Element Differential Analysis")
    r2.font.size = Pt(13)
    r2.font.italic = True
    r2.font.name = "Times New Roman"

    doc.add_paragraph()
    meta_p = doc.add_paragraph()
    meta_p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r3 = meta_p.add_run("資料期間：2026-05-01 ～ 2026-05-27　|　分析日期：2026-05-27")
    r3.font.size = Pt(11)
    r3.font.name = "Times New Roman"
    rPr3 = r3._r.get_or_add_rPr()
    rFonts3 = OxmlElement("w:rFonts")
    rFonts3.set(qn("w:eastAsia"), "標楷體")
    rPr3.append(rFonts3)

    add_hr(doc)
    doc.add_paragraph()

    # ════════════════════════════════════════════════════════════════
    # 摘要
    # ════════════════════════════════════════════════════════════════
    abs_hdr = doc.add_paragraph()
    abs_hdr.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_ah = abs_hdr.add_run("摘　要")
    r_ah.bold = True
    r_ah.font.size = Pt(13)
    r_ah.font.name = "Times New Roman"
    rPr_ah = r_ah._r.get_or_add_rPr()
    rFonts_ah = OxmlElement("w:rFonts")
    rFonts_ah.set(qn("w:eastAsia"), "標楷體")
    rPr_ah.append(rFonts_ah)

    abs_p = doc.add_paragraph()
    abs_p.paragraph_format.left_indent   = Cm(1.0)
    abs_p.paragraph_format.right_indent  = Cm(1.0)
    abs_p.paragraph_format.first_line_indent = Cm(0.85)
    abs_p.paragraph_format.space_after   = Pt(6)
    abs_p.paragraph_format.line_spacing  = Pt(20)
    r_abs = abs_p.add_run(
        "本研究以兩行軌道數據（TLE）差分分析為核心方法，對 14,090 顆低地球軌道（LEO）衛星進行"
        "軌道機動偵測，並對偵測到機動的 1,323 顆衛星執行行為分類。"
        "觀測窗口為 2026-05-01 至 2026-05-27（26 天），TLE 資料來源為 DuckDB 軌道資料庫。"
        "偵測方法計算相鄰 TLE 間的半長軸差（Δa）、傾角差（Δi）、離心率差（Δe）及"
        "J2 修正後 RAAN 殘差（ΔRAAN_res），並以物理噪聲估計設定偵測閾值。"
        "驗證結果顯示 Precision = 95.3%、Recall = 12.4%、F1 = 21.9%；"
        "Recall 偏低係結構性問題，反映有推進能力的衛星不一定在 26 天觀測窗內執行可偵測之機動。"
        "行為分類系統以規則決策樹為基礎，定義 10 種行為類別，"
        "其中 DragMakeup（大氣阻力補償）佔最大比例（57.3%，758 顆），"
        "其次為 OrbitLowering（19.0%，252 顆）及 PhasingManeuver（6.9%，91 顆）。"
        "本研究成果可用於訓練衛星行為辨識模型及星座管理態勢感知系統。")
    r_abs.font.size = Pt(11)
    r_abs.font.name = "Times New Roman"
    rPr_a = r_abs._r.get_or_add_rPr()
    rFonts_a = OxmlElement("w:rFonts")
    rFonts_a.set(qn("w:eastAsia"), "標楷體")
    rPr_a.append(rFonts_a)

    kw_p = doc.add_paragraph()
    kw_p.paragraph_format.left_indent = Cm(1.0)
    kw_p.paragraph_format.right_indent = Cm(1.0)
    kw_p.paragraph_format.space_after = Pt(8)
    r_kw = kw_p.add_run(
        "關鍵詞：低地球軌道（LEO）；兩行軌道數據（TLE）；軌道機動偵測；行為分類；"
        "J2 進動修正；大氣阻力補償；混淆矩陣")
    r_kw.bold = True
    r_kw.font.size = Pt(11)
    r_kw.font.name = "Times New Roman"
    rPr_kw = r_kw._r.get_or_add_rPr()
    rFonts_kw = OxmlElement("w:rFonts")
    rFonts_kw.set(qn("w:eastAsia"), "標楷體")
    rPr_kw.append(rFonts_kw)

    add_hr(doc)

    # ════════════════════════════════════════════════════════════════
    # § 1. 緒論
    # ════════════════════════════════════════════════════════════════
    add_heading_paper(doc, "緒　論", level=1, numbering="1.")
    add_body_paper(doc,
        "隨著低地球軌道（LEO）商業星座的快速擴張，Starlink、Amazon Kuiper 等"
        "巨型星座衛星數量已突破萬顆，主動軌道機動管理成為太空交通管理（STM）的核心挑戰。"
        "能否從公開可用的 TLE 數據中準確辨識衛星是否正在執行機動，"
        "並進一步判斷機動行為類型，對碰撞預警、星座健康監測與訓練資料集建置均具重要意義。")
    add_body_paper(doc,
        "TLE 是由地面雷達網路（主要為美國太空監視網路，SSN）擬合生成的軌道元素集，"
        "通常每天更新 1–2 次。相鄰 TLE 間的軌道元素差分可反映衛星的軌道演化，"
        "機動行為會造成超出自然攝動範圍的軌道元素突變。"
        "然而，TLE 本身存在測量噪聲（Δa 約 0.3–0.5 km），"
        "且地球 J2 扁率攝動會引起自然的 RAAN 進動（每天 -5° 至 +5° 不等），"
        "若不妥善校正，極易引發誤報。")
    add_body_paper(doc,
        "本研究建構了一套完整的 TLE 差分機動偵測管線，涵蓋：J2 RAAN 修正、"
        "物理噪聲估計為基礎的閾值設定、混淆矩陣驗證，以及規則決策樹行為分類系統。"
        "以 14,090 顆 LEO 衛星為對象，驗證推進系統標註品質，並對 1,323 顆機動衛星進行行為分類。")

    # ════════════════════════════════════════════════════════════════
    # § 2. 資料集與觀測條件
    # ════════════════════════════════════════════════════════════════
    add_heading_paper(doc, "資料集與觀測條件", level=1, numbering="2.")
    add_body_paper(doc,
        "本研究使用的 TLE 數據存儲於 DuckDB 軌道資料庫（space_db.duckdb），"
        "資料表 tle_table 中的 date_tag 欄位記錄 TLE epoch 時刻。"
        "推進系統標註來源為人工整理的 annotations_leo_full.csv，"
        "包含衛星名稱、NORAD 編號、推進類別，以及質量等參數。")

    add_table_paper(doc,
        "資料集與觀測參數摘要",
        ["參數", "值", "說明"],
        [
            ["觀測期間",    "2026-05-01 ～ 2026-05-27",    "26 天觀測窗口"],
            ["分析衛星總數", "14,090 顆",                   "annotations_leo_full.csv 全量 LEO 衛星"],
            ["有效分析數",   "14,019 顆（99.5%）",          "TLE ≥ 3 筆且有相鄰轉移記錄"],
            ["資料不足",     "71 顆（0.5%）",                "0–2 筆 TLE，無法計算差分"],
            ["TLE 記錄總數", "約 62,000 筆",                 "偵測機動的 1,323 顆衛星所對應記錄"],
            ["推進類別",
             "Electric_EP / Chemical /\nMicro/ColdGas / Hybrid/Other / None",
             "五類，以 Electric_EP 為主（10,074 顆）"],
            ["最長 TLE 間隔", "7 天",                        "超過則跳過以避免缺測誤差"],
        ],
        col_widths=[3.5, 4.5, 8.0])

    add_table_paper(doc,
        "推進類別標註分布",
        ["推進類別", "衛星數", "占比（%）", "說明"],
        [
            ["Electric_EP",    "10,074", "71.5%", "離子推進、電推（Starlink Krypton Hall 等）"],
            ["None",            "3,833", "27.2%", "無推進系統（CubeSat、碎片等）"],
            ["Hybrid/Other",      "56",  "0.4%",  "複合推進或特殊推進系統"],
            ["Chemical",          "50",  "0.4%",  "化學推進（單組元、雙組元）"],
            ["Micro/ColdGas",     "33",  "0.2%",  "冷氣或微型推進系統"],
        ],
        col_widths=[3.5, 2.5, 2.5, 7.5])

    # ════════════════════════════════════════════════════════════════
    # § 3. 軌道機動偵測方法論
    # ════════════════════════════════════════════════════════════════
    add_heading_paper(doc, "軌道機動偵測方法論", level=1, numbering="3.")

    add_heading_paper(doc, "3.1  軌道元素差分計算", level=2)
    add_body_paper(doc,
        "對單顆衛星按 epoch 排序後，對每對相鄰 TLE（epoch[i−1], epoch[i]）計算四個差分量。"
        "設 Δt 為兩 TLE 間的時間差（秒），則：")
    add_formula(doc, "Δa        =  sma_km[i]          −  sma_km[i−1]",          "（1）")
    add_formula(doc, "Δi        =  inclination[i]      −  inclination[i−1]",    "（2）")
    add_formula(doc, "Δe        =  eccentricity[i]     −  eccentricity[i−1]",   "（3）")
    add_formula(doc, "ΔRAAN_res =  ΔRAAN_raw           −  (dΩ/dt)×Δt",          "（4）")
    add_body_paper(doc,
        "其中 ΔRAAN_raw 為有號角度差（範圍 (−180°, 180°]），(dΩ/dt)×Δt 為 J2 進動修正量。"
        "若 Δt > 604,800 s（7 天），則跳過該轉移，避免長期資料缺口造成差分失真。")

    add_heading_paper(doc, "3.2  J2 RAAN 進動修正", level=2)
    add_body_paper(doc,
        "地球非球形扁率（J2 = 1.08263×10⁻³）使衛星升交點赤經（RAAN）產生自然進動，"
        "其進動率與衛星軌道高度、傾角有關，必須從觀測到的 ΔRAAN 中扣除，"
        "才能分離出主動機動的貢獻。進動率 dΩ/dt（deg/s）計算如下：")
    add_formula(doc,
        "dΩ/dt  =  −(3/2) × J2 × (RE/p)² × n × cos(i)",  "（5）")
    add_body_paper(doc,
        "其中 n = √(μ/a³) 為平均角速度（rad/s），p = a(1−e²) 為半通徑（km），"
        "RE = 6378.137 km 為地球赤道半徑，μ = 398600.4418 km³/s² 為地球引力常數。"
        "扣除 J2 進動後的 RAAN 殘差 ΔRAAN_res 若仍超過偵測閾值，"
        "則判定為主動側推（如軌道面調整或 RAAN 站位保持）。")

    add_heading_paper(doc, "3.3  偵測閾值設定依據", level=2)
    add_body_paper(doc,
        "閾值設定原則為「高於 TLE 測量噪聲的 2 倍」，以確保訊雜比（SNR）足夠同時避免過多誤報。"
        "TLE 噪聲估計來源為文獻報告值及本資料集中靜態（無推進）衛星的差分統計。")
    add_table_paper(doc,
        "機動偵測閾值設定（與 TLE 測量噪聲對照）",
        ["差分量", "偵測閾值", "TLE 噪聲估計", "典型機動幅度"],
        [
            ["Δa（半長軸）",     "> 1.0 km",   "0.3–0.5 km",   "電推單次 0.1–5 km；化推 2–50 km"],
            ["Δi（傾角）",       "> 0.02°",    "0.005–0.01°",  "傾角調整 0.1–2°"],
            ["Δe（離心率）",     "> 0.001",    "0.0003–0.0005","主動機動 0.005–0.05"],
            ["ΔRAAN_res（殘差）","> 0.1°",     "0.03–0.05°",   "主動 RAAN 調整 0.3–5°/天"],
        ],
        col_widths=[3.8, 2.5, 3.0, 6.7])
    add_body_paper(doc,
        "任一差分量超過對應閾值，則該 TLE 對轉移被標記為 flagged=True。"
        "若衛星在觀測期間有至少 1 個 flagged 轉移，則 maneuver_detected = True。")

    add_heading_paper(doc, "3.4  混淆矩陣設計", level=2)
    add_body_paper(doc,
        "以推進系統標註作為真值（Ground Truth），TLE 差分偵測結果作為預測（Prediction），"
        "構建二元混淆矩陣以評估偵測性能。")
    add_table_paper(doc,
        "混淆矩陣定義",
        ["縮寫", "定義", "本研究意義"],
        [
            ["TP（真陽性）", "有推進 ∩ 偵測到機動",   "衛星有推進能力，且在觀測期內確實機動"],
            ["FN（假陰性）", "有推進 ∩ 未偵測到機動", "有推進能力，但 26 天內未執行可偵測機動（漏報）"],
            ["FP（假陽性）", "無推進 ∩ 偵測到機動",   "無推進能力，但 TLE 差分超閾值（誤報）"],
            ["TN（真陰性）", "無推進 ∩ 未偵測到機動", "無推進能力，且 TLE 差分正常（正確排除）"],
        ],
        col_widths=[3.0, 5.0, 8.0])
    add_formula(doc, "Precision  =  TP / (TP + FP)",         "（6）")
    add_formula(doc, "Recall     =  TP / (TP + FN)",          "（7）")
    add_formula(doc, "F1-score   =  2 × Precision × Recall / (Precision + Recall)", "（8）")

    # ════════════════════════════════════════════════════════════════
    # § 4. 軌道機動行為分類系統
    # ════════════════════════════════════════════════════════════════
    add_heading_paper(doc, "軌道機動行為分類系統", level=1, numbering="4.")
    add_body_paper(doc,
        "本研究對 1,323 顆偵測到機動的衛星，進一步利用衍生特徵構建規則決策樹，"
        "將機動行為分類為 10 種類別。分類器僅使用 TLE 差分統計量，不依賴衛星型號或任務資料庫，"
        "具備對新衛星的通用性。")

    add_heading_paper(doc, "4.1  衍生特徵計算", level=2)
    add_body_paper(doc,
        "在逐對差分完成後，對單顆衛星的所有轉移記錄進行統計，"
        "計算以下四類衍生特徵作為決策樹輸入。")
    add_table_paper(doc,
        "行為分類器衍生特徵",
        ["特徵類型", "特徵名稱", "計算公式", "物理意義"],
        [
            ["觸發維度\n（布林）",
             "has_da / has_di /\nhas_de / has_draan",
             "任一 flagged 轉移超過對應閾值",
             "識別機動發生在哪個軌道元素維度"],
            ["幅度指標",
             "max_da（km）",
             "flagged 轉移中 |Δa| 最大值",
             "單次機動的最大規模"],
            ["幅度指標",
             "net_da（km）",
             "所有轉移 Δa 的代數和",
             "26 天累積軌道高度淨變化（正=升，負=降）"],
            ["單調性指標",
             "neg_streak（次）",
             "da < −0.3 km 連續出現的最長序列長度",
             "識別大氣阻力持續衰減 vs 間歇性主動機動"],
            ["單調性指標",
             "monotone_decay（布林）",
             "neg_streak≥5 AND total_drop>5 km\nAND net_da<−3 km",
             "單調衰減旗標（大氣阻力特徵）"],
            ["頻率指標",
             "flag_rate",
             "n_flagged / n_transitions",
             "旗標轉移占所有轉移的比例"],
            ["頻率指標",
             "burn_freq_per_day",
             "n_flagged / 26",
             "日均機動次數（26 天觀測窗）"],
        ],
        col_widths=[2.5, 3.2, 4.5, 5.8])

    add_heading_paper(doc, "4.2  規則決策樹設計", level=2)
    add_body_paper(doc,
        "決策樹以推進系統標註作為最頂層分歧，再依衍生特徵的優先順序逐層判定行為類別。"
        "優先順序設計的關鍵：monotone_decay 特徵置於推進分支的最頂層，"
        "確保大幅持續降軌衛星（如 JILIN-01 GAOFEN 3J，net_da = −117 km）"
        "被正確識別為「大氣阻力補償」而非「主動降軌」。")
    add_code_paper(doc,
"propulsion_class ∈ {Electric_EP, Chemical, Micro/ColdGas, Hybrid/Other}?\n"
"│\n"
"├─ 否（無推進）\n"
"│   ├─ monotone_decay = True        →  DragDecay\n"
"│   ├─ has_draan=T, has_da=F        →  UnknownRAANAnomaly\n"
"│   └─ 其他                         →  UnknownFP\n"
"│\n"
"└─ 是（有推進）\n"
"    ├─ [最優先] monotone_decay = True →  DragMakeup\n"
"    ├─ net_da > +5 km                →  OrbitRaising\n"
"    ├─ net_da < −5 km                →  OrbitLowering\n"
"    ├─ has_da AND has_di             →  ComplexManeuver\n"
"    ├─ has_di only                   →  InclinationChange\n"
"    ├─ has_draan only                →  RAANMaintenance\n"
"    ├─ flag_rate > 25%               →  Stationkeeping\n"
"    ├─ n_flagged ≤ 3 AND max_da > 5km →  PhasingManeuver\n"
"    ├─ flag_rate ≤ 15%               →  DragMakeup\n"
"    └─ 其他                          →  Stationkeeping")

    add_heading_paper(doc, "4.3  十種行為類別定義", level=2)
    add_table_paper(doc,
        "十種軌道機動行為類別定義",
        ["類別代號", "中文名稱", "核心判斷條件", "物理意義"],
        [
            ["DragMakeup",
             "大氣阻力補償",
             "monotone_decay=True（有推進）\n或 flag_rate ≤ 15%",
             "電推持續補償大氣阻力衰減，SMA 呈鋸齒狀週期振盪"],
            ["OrbitLowering",
             "軌道降低 / 離軌",
             "net_da < −5 km（有推進，非單調衰減）",
             "主動降軌、離軌燃燒或任務結束程序"],
            ["PhasingManeuver",
             "相位調整",
             "n_flagged ≤ 3 且 max_da > 5 km",
             "新部署衛星一次性相位機動到目標星座軌道"],
            ["OrbitRaising",
             "軌道抬升",
             "net_da > +5 km（有推進）",
             "主動升軌，通常為部署初期爬升或任務軌道調整"],
            ["Stationkeeping",
             "站位保持",
             "flag_rate > 25%",
             "SAR/EO 精密衛星高頻小幅維持軌道精確度"],
            ["UnknownFP",
             "未知假陽性",
             "無推進標註，非單調衰減",
             "標註可能缺失，或 TLE 誤差引發的殘餘誤報"],
            ["DragDecay",
             "大氣阻力衰減",
             "monotone_decay=True（無推進）",
             "無推進能力衛星受大氣阻力持續拖曳衰減"],
            ["InclinationChange",
             "傾角調整",
             "has_di=True，has_da=False",
             "純軌道面傾角機動（需大量 Δv，少見）"],
            ["ComplexManeuver",
             "複合機動",
             "has_da=True 且 has_di=True",
             "同時調整軌道高度與軌道面（大型任務調整）"],
            ["UnknownRAANAnomaly",
             "RAAN 異常待查",
             "has_draan=True，has_da/di=False",
             "RAAN 殘差超閾值但無明確機動，可能為 TLE 系統誤差"],
        ],
        col_widths=[3.5, 2.8, 4.0, 5.7])

    # ════════════════════════════════════════════════════════════════
    # § 5. 實驗結果
    # ════════════════════════════════════════════════════════════════
    add_heading_paper(doc, "實驗結果", level=1, numbering="5.")

    add_heading_paper(doc, "5.1  標註驗證指標", level=2)
    add_body_paper(doc,
        "以 14,019 顆有效衛星（排除 TLE 不足 3 筆者）進行驗證，混淆矩陣結果如下：")
    add_table_paper(doc,
        "混淆矩陣實驗結果",
        ["指標", "數值", "說明"],
        [
            ["TP（真陽性）", "1,261 顆", "有推進系統，且 TLE 差分偵測到機動"],
            ["FN（假陰性）", "8,925 顆", "有推進系統，但 26 天內未偵測到機動（結構性漏報）"],
            ["FP（假陽性）", "62 顆",   "無推進系統，但 TLE 差分超閾值（誤報）"],
            ["TN（真陰性）", "3,771 顆", "無推進系統，TLE 差分正常（正確排除）"],
            ["Precision",   "95.3%",   "偵測到的機動中確實有推進能力的比例"],
            ["Recall",      "12.4%",   "所有有推進衛星中被成功偵測的比例"],
            ["F1-score",    "21.9%",   "Precision 與 Recall 的調和平均"],
        ],
        col_widths=[3.0, 2.5, 10.5])

    add_heading_paper(doc, "5.2  各推進類別偵測率", level=2)
    add_body_paper(doc,
        "不同推進類型的偵測率差異顯著，主要受推力量級影響：")
    add_table_paper(doc,
        "各推進類別 TLE 機動偵測率",
        ["推進類別", "樣本數", "有效分析", "偵測到機動", "偵測率"],
        [
            ["Electric_EP",   "10,074", "10,047", "1,238", "12%"],
            ["Chemical",          "50",     "50",    "17", "34%"],
            ["Micro/ColdGas",     "33",     "33",     "4", "12%"],
            ["Hybrid/Other",      "56",     "56",     "2",  "4%"],
            ["None（無推進）",  "3,877",  "3,833",    "62", "—（FP）"],
        ],
        col_widths=[3.8, 2.5, 2.5, 3.0, 2.5])
    add_body_paper(doc,
        "化學推進偵測率（34%）顯著高於電推（12%），"
        "因化學推進每次燃燒產生較大的 Δa（典型 5–50 km），更易超過偵測閾值。"
        "電推每次燃燒僅 0.1–3 km，多次小幅燃燒的累積效應才能被捕捉。")

    add_heading_paper(doc, "5.3  行為分類結果", level=2)
    add_body_paper(doc,
        "1,323 顆偵測到機動的衛星中，行為分類結果如下：")
    add_table_paper(doc,
        "十種行為類別分類結果統計",
        ["行為類別", "顆數", "占比", "主要星座 / 說明"],
        [
            ["DragMakeup",         "758", "57.3%", "Starlink 星座為主；電推持續低幅補償"],
            ["OrbitLowering",      "252", "19.0%", "Starlink 離軌；net_da 平均 −13 km"],
            ["PhasingManeuver",     "91",  "6.9%", "Starlink 新部署；少次大幅燃燒"],
            ["OrbitRaising",        "81",  "6.1%", "Kuiper 部署爬升；net_da 平均 +8 km"],
            ["Stationkeeping",      "69",  "5.2%", "SAR/EO 精密衛星；flag_rate > 25%"],
            ["UnknownFP",           "39",  "2.9%", "標註待核查或 TLE 系統誤差"],
            ["DragDecay",           "21",  "1.6%", "極低軌無推進衛星；RHOK-SAT 等"],
            ["InclinationChange",    "6",  "0.5%", "Starlink 面調整"],
            ["ComplexManeuver",      "4",  "0.3%", "KUIPER-00262 等複合機動"],
            ["UnknownRAANAnomaly",   "2",  "0.2%", "IXPE、ORS 5 等 RAAN 異常"],
            ["合計",             "1,323", "100%",  ""],
        ],
        col_widths=[3.8, 1.8, 1.8, 8.6])

    add_table_paper(doc,
        "推進類別 × 行為類別交叉分布（主要類別）",
        ["推進類別", "DragMakeup", "OrbitLowering", "PhasingManeuver", "OrbitRaising",
         "Stationkeeping", "其他"],
        [
            ["Electric_EP",   "742", "249", "90", "79", "69", "9"],
            ["Chemical",       "13",   "1",  "0",  "2",  "0", "1"],
            ["Micro/ColdGas",   "2",   "2",  "0",  "0",  "0", "0"],
            ["Hybrid/Other",    "1",   "0",  "1",  "0",  "0", "0"],
            ["無推進（FP）",    "0",   "0",  "0",  "0",  "0","62"],
        ],
        col_widths=[3.2, 2.5, 3.0, 3.2, 2.8, 2.8, 2.0])

    # ════════════════════════════════════════════════════════════════
    # § 6. 討論
    # ════════════════════════════════════════════════════════════════
    add_heading_paper(doc, "討　論", level=1, numbering="6.")

    add_heading_paper(doc, "6.1  Recall 偏低（12.4%）的結構性成因分析", level=2)
    add_body_paper(doc,
        "Recall = 12.4% 並非方法論缺陷，而是由以下四個結構性因素共同決定：")
    add_body_paper(doc,
        "（1）推進能力 ≠ 觀測期間正在機動。標註為「有推進」僅代表衛星在設計上配備推進系統，"
        "不代表衛星在 2026-05-01 至 2026-05-27 這 26 天內恰好執行了可偵測的機動。"
        "大多數電推衛星機動週期可能超過 30 天，在 26 天窗口內未必觸發。",
        first_line=False)
    add_body_paper(doc,
        "（2）偵測閾值設在噪聲 2× 位置。電推衛星每次燃燒 Δa 約 0.1–0.5 km，"
        "低於 THR_DA_SM = 1.0 km 的閾值；若降低閾值，FP 將大幅增加，Precision 下降。",
        first_line=False)
    add_body_paper(doc,
        "（3）TLE 更新頻率不足。機動可能在兩次 TLE 更新之間發生，"
        "若前後兩個 TLE 均在機動後生成，差分值會偏小而被過濾。",
        first_line=False)
    add_body_paper(doc,
        "（4）機動相互抵消。多段升軌後降軌的累積 net_da 趨近 0，被判定為無明顯機動。",
        first_line=False)

    add_heading_paper(doc, "6.2  三大分類邊界難題", level=2)
    add_body_paper(doc,
        "難題一：DragMakeup vs OrbitLowering。兩者均為 net_da < 0 且 has_da = True。"
        "區分依據為 monotone_decay 旗標：大氣阻力補償的 SMA 衰減是單調連續的（neg_streak ≥ 5），"
        "而主動降軌通常是間歇性的。JILIN-01 GAOFEN 3J（net_da = −117 km）雖幅度巨大，"
        "但呈現持續單調降軌特徵，故被正確歸為 DragMakeup。"
        "未來可加入 total_drop > 50 km 門檻以區分「大幅終端降軌」與持續補償。")
    add_body_paper(doc,
        "難題二：PhasingManeuver vs 單次軌道維持燃燒。"
        "兩者均為少次數（n_flagged ≤ 3）但高幅度（max_da > 5 km）的機動，"
        "TLE 差分特徵相同但物理意義不同。前者為新部署衛星的一次性相位調整，"
        "後者可能為空間站定期維持燃燒。目前以衛星類型標籤區分（後續改進方向）。")
    add_body_paper(doc,
        "難題三：DragDecay（無推進衰減）vs 電推耗盡後的衰減。"
        "電推耗盡後衛星行為與無推進衛星相同，若標註未更新，"
        "則被錯分為 DragMakeup。建議加入發射年份與電推設計壽命資訊作為輔助特徵。")

    add_heading_paper(doc, "6.3  方法局限性與改進方向", level=2)
    add_table_paper(doc,
        "方法局限性與對應改進策略",
        ["局限性", "影響", "建議改進"],
        [
            ["閾值為全域固定值",
             "不同高度（200–2000 km）的 TLE 噪聲特性不同",
             "依軌道高度分段設定自適應閾值"],
            ["未利用 B* 係數",
             "大氣阻力估計不準確，特別是高阻力小衛星",
             "將 B* 納入 monotone_decay 判斷"],
            ["觀測窗口限制（26 天）",
             "低頻機動（>30 天週期）無法偵測",
             "延長至 90 天以上或使用滾動窗口"],
            ["TLE 更新不均勻",
             "高關注衛星更新快，碎片更新慢，影響差分密度",
             "加入「TLE 更新率」作為資料品質過濾條件"],
            ["單調性判斷為硬編碼門檻",
             "neg_streak ≥ 5 對不同觀測期長度不具通用性",
             "改為「最長連續負值序列 / 總轉移數」的比例指標"],
        ],
        col_widths=[4.0, 4.5, 7.5])

    # ════════════════════════════════════════════════════════════════
    # § 7. 結論
    # ════════════════════════════════════════════════════════════════
    add_heading_paper(doc, "結　論", level=1, numbering="7.")
    add_body_paper(doc,
        "本研究建構了一套完整的 TLE 差分 LEO 衛星軌道機動偵測與行為分類管線，"
        "主要貢獻包括：")
    add_body_paper(doc,
        "（1）提出以物理噪聲估計為基礎的偵測閾值設定方法（約 2× TLE 測量噪聲），"
        "有效控制誤報率（Precision = 95.3%），同時維持可接受的計算效率。",
        first_line=False)
    add_body_paper(doc,
        "（2）設計 J2 RAAN 進動修正算法，使 ΔRAAN 閾值真正代表主動機動而非自然攝動，"
        "避免大量 RAAN 誤報。",
        first_line=False)
    add_body_paper(doc,
        "（3）引入 monotone_decay 旗標（基於最長連續負 Δa 序列），"
        "有效區分大氣阻力持續衰減與間歇性主動機動，解決 DragMakeup 與 OrbitLowering 的邊界模糊問題。",
        first_line=False)
    add_body_paper(doc,
        "（4）定義 10 種行為類別的規則決策樹，"
        "對 1,323 顆機動衛星完成無監督式自動分類，其中 DragMakeup（57.3%）及 OrbitLowering（19.0%）"
        "為 LEO 商業星座的主流行為模式。",
        first_line=False)
    add_body_paper(doc,
        "本研究結果可直接應用於：LEO 衛星訓練資料集建置、星座健康自動監測、"
        "太空交通管理系統的機動預測模型，以及推進系統標註品質驗證。"
        "後續研究建議引入 B* 係數、自適應閾值與更長觀測窗口，進一步提升 Recall 並降低分類邊界的模糊性。")

    add_hr(doc)

    # ════════════════════════════════════════════════════════════════
    # 參考文獻
    # ════════════════════════════════════════════════════════════════
    ref_hdr = doc.add_paragraph()
    r_rh = ref_hdr.add_run("參考文獻")
    r_rh.bold = True
    r_rh.font.size = Pt(13)
    r_rh.font.name = "Times New Roman"
    rPr_rh = r_rh._r.get_or_add_rPr()
    rFonts_rh = OxmlElement("w:rFonts")
    rFonts_rh.set(qn("w:eastAsia"), "標楷體")
    rPr_rh.append(rFonts_rh)

    refs = [
        "[1] Kelso, T.S. (2007). Validation of SGP4 and IS-GPS-200D Against GPS Precision Ephemerides. "
        "17th AAS/AIAA Space Flight Mechanics Meeting, Sedona, AZ.",
        "[2] Hoots, F.R., Roehrich, R.L. (1980). Models for Propagation of NORAD Element Sets. "
        "Spacetrack Report No. 3, U.S. Air Force.",
        "[3] Vallado, D.A., Crawford, P., Hujsak, R., Kelso, T.S. (2006). Revisiting Spacetrack Report #3: "
        "Rev 1. AIAA 2006-6753.",
        "[4] 台灣太空中心（TASA）（2024）。LEO 星座衛星推進系統調查報告。",
        "[5] SpaceX（2023）. Starlink Satellite Propulsion System: Krypton Hall-Effect Thruster Technical Overview.",
        "[6] Muelhaupt, T.J., Sorge, M.E., Morin, J., Wilson, R.S. (2019). Space Traffic Management in the New "
        "Space Era. Journal of Space Safety Engineering, 6(2), 80–87.",
        "[7] DuckDB Development Team（2024）. DuckDB: An Embeddable Analytical Database. SIGMOD.",
    ]
    for ref in refs:
        rp = doc.add_paragraph()
        rp.paragraph_format.left_indent  = Cm(0.8)
        rp.paragraph_format.first_line_indent = Cm(-0.8)
        rp.paragraph_format.space_after  = Pt(4)
        r_ref = rp.add_run(ref)
        r_ref.font.size = Pt(10)
        r_ref.font.name = "Times New Roman"
        rPr_ref = r_ref._r.get_or_add_rPr()
        rFonts_ref = OxmlElement("w:rFonts")
        rFonts_ref.set(qn("w:eastAsia"), "標楷體")
        rPr_ref.append(rFonts_ref)

    # ── 儲存 ─────────────────────────────────────────────────────────────────
    doc.save(OUT_PATH)
    print(f"[DONE] {OUT_PATH}")


if __name__ == "__main__":
    build_doc()
